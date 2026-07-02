#!/usr/bin/env python3
"""Generate SWP Knowledge Assessment Document"""
import subprocess, sys
subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'], check=True)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = '/Users/sandeeppr/CIBOP/PPT 07 Derivative SWAPS/SWP_Knowledge_Assessment_v2.docx'

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def header_para(doc, text, color_hex=(0x1E, 0x3A, 0x5F), size=12, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color_hex)
    return p

def add_q_header(doc, q_num, q_type, ear, uor_id, video_ref):
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    widths = [0.6, 1.8, 1.8, 3.1]
    vals = [q_num, q_type, f'EAR: {ear}', f'{uor_id} | {video_ref}']
    fills = ['1E3A5F', '005F5F', '8B6914', 'EEEEEE']
    text_colors = [(255,255,255),(255,255,255),(255,255,255),(30,30,30)]
    for i, (val, w, fill, tc_col) in enumerate(zip(vals, widths, fills, text_colors)):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, fill)
        set_col_width(cell, w)
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(str(val))
        run2.bold = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(*tc_col)
    doc.add_paragraph()

def two_col_table(doc, col1_header, col1_data, col2_header, col2_data, fill1='EEF4FB', fill2='F0F7F4'):
    t = doc.add_table(rows=1+max(len(col1_data), len(col2_data)), cols=2)
    t.style = 'Table Grid'
    # Header row
    for i, (hdr, fill) in enumerate(zip([col1_header, col2_header], [fill1, fill2])):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '1E3A5F')
        set_col_width(cell, 3.5)
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(hdr)
        run2.bold = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(255,255,255)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in range(max(len(col1_data), len(col2_data))):
        row = t.rows[r+1]
        v1 = col1_data[r] if r < len(col1_data) else ''
        v2 = col2_data[r] if r < len(col2_data) else ''
        for i, (v, fill) in enumerate(zip([v1, v2], [fill1, fill2])):
            cell = row.cells[i]
            set_cell_bg(cell, fill)
            set_col_width(cell, 3.5)
            p2 = cell.paragraphs[0]
            run2 = p2.add_run(v)
            run2.font.size = Pt(9)
    doc.add_paragraph()

def answer_box(doc, correct_ans, reasoning):
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'
    labels = ['CORRECT ANSWER', 'REASONING']
    vals = [correct_ans, reasoning]
    fills = ['E8F5E9', 'FFF8E1']
    for r, (lbl, val, fill) in enumerate(zip(labels, vals, fills)):
        cell0 = t.rows[r].cells[0]
        set_cell_bg(cell0, '1E3A5F')
        set_col_width(cell0, 1.5)
        p2 = cell0.paragraphs[0]
        run2 = p2.add_run(lbl)
        run2.bold = True
        run2.font.size = Pt(8.5)
        run2.font.color.rgb = RGBColor(255,255,255)

        cell1 = t.rows[r].cells[1]
        set_cell_bg(cell1, fill)
        set_col_width(cell1, 5.8)
        p3 = cell1.paragraphs[0]
        run3 = p3.add_run(val)
        run3.font.size = Pt(9)
    doc.add_paragraph()

# ─────────────────────────────────────────────
# QUESTION BANK — 4 questions per UOR × 6 UORs = 24 questions
# Types cycle: MCQ Inline, True/False Multi, Gap Fill Dropdown, Matching Pairs,
#              Table Selection (then back to MCQ for 5th cycle)
# ─────────────────────────────────────────────

QUESTIONS = {
    'SWP.1': [
        {
            'num': 'QK1', 'type': 'MCQ INLINE', 'ear': 'DEFINE',
            'video_ref': 'SWP.1 Video 1',
            'q': 'In SWP.1 Video 1, Ryo explains the fundamental nature of a swap. A swap is BEST described as:',
            'options': [
                'A) An exchange-traded contract obligating the buyer to purchase an asset at a future date',
                'B) An OTC agreement between two parties to exchange a series of cash flows based on different terms over an agreed period',
                'C) A one-time payment from the buyer to the seller in exchange for an asset',
                'D) A government-backed contract to fix interest rates on sovereign bonds',
            ],
            'correct': 'B',
            'correct_text': 'B) An OTC agreement between two parties to exchange a series of cash flows based on different terms',
            'reasoning': 'From Video 1: "A swap is simply an agreement to exchange those cash flow obligations — no loan transfer, no asset movement. Just a contractual exchange of payment streams." Swaps are OTC derivatives — bilateral, customisable, not exchange-traded.',
        },
        {
            'num': 'QK2', 'type': 'TRUE/FALSE MULTI', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.1 Video 1 & 2',
            'q': 'Based on SWP.1 Videos 1 and 2, mark each statement True or False.',
            'statements': [
                ('In a swap, the two parties are called the fixed payer and the floating payer', True),
                ('Swaps are exchange-traded derivatives, similar to futures contracts', False),
                ('The notional principal in a swap is actually transferred between parties at inception', False),
                ('Swaps can be used for hedging, speculation, and exploiting comparative advantage', True),
                ('An ISDA Master Agreement typically governs the legal terms of an OTC swap', True),
                ('A swap is economically equivalent to a series of forward rate agreements', True),
            ],
            'correct_true': [1, 4, 5, 6],
            'correct_false': [2, 3],
            'correct_text': 'True: 1, 4, 5, 6 | False: 2, 3',
            'reasoning': 'From Video 1 & 2: Swaps are OTC (not exchange-traded). The notional is NEVER exchanged — only net interest differences. The two parties are fixed payer and floating payer. ISDA Master Agreement governs OTC derivatives. A swap can be decomposed into a series of forwards.',
        },
        {
            'num': 'QK3', 'type': 'GAP FILL DROPDOWN', 'ear': 'EXPLAIN',
            'video_ref': 'SWP.1 Video 2',
            'q': 'SWP.1 Video 2 explains the comparative advantage argument for swaps. Complete the explanation.',
            'paragraph': 'Company A can borrow fixed at _1_ or floating at MIBOR+1%. Company B can borrow fixed at 10% or floating at _2_. Company A has comparative advantage in _3_ markets; Company B in floating markets. If each borrows where they have the advantage and then _4_, both achieve their preferred rate at a lower all-in cost. The total saving available to share between them equals the _5_ in their borrowing rate differentials across the two markets. This is an arbitrage of _6_.',
            'gaps': [
                ('_1_', ['8%', '6%', '12%', '5%'], '8%'),
                ('_2_', ['MIBOR+2%', 'MIBOR+0.5%', 'MIBOR+3%', '9% fixed'], 'MIBOR+2%'),
                ('_3_', ['fixed', 'floating', 'equity', 'commodity'], 'fixed'),
                ('_4_', ['swap obligations', 'sell their loans', 'default', 'merge companies'], 'swap obligations'),
                ('_5_', ['difference', 'sum', 'product', 'average'], 'difference'),
                ('_6_', ['credit spreads across markets', 'equity valuations', 'FX rates', 'regulatory capital requirements'], 'credit spreads across markets'),
            ],
            'correct_text': '_1_ 8% | _2_ MIBOR+2% | _3_ fixed | _4_ swap obligations | _5_ difference | _6_ credit spreads across markets',
            'reasoning': 'From Video 2: "Infra Co can borrow fixed at 8% or floating at MIBOR+1%. Tech Co can borrow fixed at 10% or floating at MIBOR+2%. Each borrows where they are cheapest, then swaps. Both end up with their preferred rate at a lower total cost. That is arbitrage of credit spreads."',
        },
        {
            'num': 'QK4', 'type': 'MATCHING PAIRS', 'ear': 'DIFFERENTIATE',
            'video_ref': 'SWP.1 Videos 1–2',
            'q': 'Match each swap concept to its correct description.',
            'col_a': ['Fixed Payer', 'Floating Payer', 'Notional Principal', 'OTC Derivative', 'Comparative Advantage Use', 'Hedging Use'],
            'col_b': [
                'Pays a set percentage rate each period regardless of market movements',
                'Pays a rate that resets based on MIBOR or similar benchmark each period',
                'Reference amount used to calculate interest — never physically exchanged',
                'Bilateral contract negotiated directly, governed by ISDA, not exchange-traded',
                'Each party borrows where relatively cheaper, then swaps to desired rate',
                'Corporate converts floating rate debt to fixed to lock in predictable cost',
            ],
            'correct_text': 'Fixed Payer → Pays set % each period | Floating Payer → pays MIBOR-based rate | Notional → reference amount never exchanged | OTC → bilateral ISDA contract | Comparative Advantage → each borrows where cheaper then swaps | Hedging → convert floating to fixed',
            'reasoning': 'From Videos 1 and 2: The three motivations for swaps are hedging, speculation, and comparative advantage. The notional is never exchanged. Fixed payer is locked at agreed rate; floating payer resets each period. OTC means bilateral and not exchange-traded.',
        },
    ],
    'SWP.2': [
        {
            'num': 'QK1', 'type': 'MCQ INLINE', 'ear': 'DESCRIBE',
            'video_ref': 'SWP.2 Video 1',
            'q': 'In SWP.2 Video 1, Aria explains the plain vanilla IRS. Company A pays 7% fixed on Rs 100 crore semi-annually; Company B pays MIBOR (currently 6%) semi-annually. What is the net settlement at this payment date?',
            'options': [
                'A) Company B pays Rs 7 crore to Company A',
                'B) Company A pays Rs 0.5 crore to Company B — because 7% fixed exceeds 6% floating on semi-annual basis',
                'C) No payment — swaps net to zero every period',
                'D) Company A pays Rs 7 crore; Company B pays Rs 6 crore — both amounts settled separately',
            ],
            'correct': 'B',
            'correct_text': 'B) Company A pays Rs 0.5 crore to Company B (7% × 100Cr × 0.5 = Rs 3.5 Cr less 6% × 100Cr × 0.5 = Rs 3 Cr → net Rs 0.5 Cr)',
            'reasoning': 'From Video 1: "Fixed payer owes: 7% × 100 crore × 0.5 = Rs 3.5 crore. Floating payer owes: 6% × 100 crore × 0.5 = Rs 3 crore. Net settlement: fixed payer pays Rs 0.5 crore." Only the net difference changes hands — not both gross payments.',
        },
        {
            'num': 'QK2', 'type': 'TRUE/FALSE MULTI', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.2 Video 1',
            'q': 'Based on SWP.2 Video 1, mark each statement about IRS and MIBOR True or False.',
            'statements': [
                ('MIBOR stands for Mumbai Interbank Offer Rate and reflects short-term bank lending rates in India', True),
                ('In an IRS, the notional principal is exchanged between parties at both inception and maturity', False),
                ('If MIBOR rises above the fixed rate, the fixed payer receives more from the floating leg than they pay on the fixed leg', True),
                ('The floating rate in an IRS resets every period based on prevailing MIBOR or LIBOR', True),
                ('A plain vanilla IRS involves three currencies and two notional amounts', False),
                ('Net settlement means only the difference between fixed and floating payments is exchanged', True),
            ],
            'correct_true': [1, 3, 4, 6],
            'correct_false': [2, 5],
            'correct_text': 'True: 1, 3, 4, 6 | False: 2, 5',
            'reasoning': 'From Video 1: Notional is NEVER exchanged in an IRS (only currency swaps exchange principal). MIBOR is the Indian floating benchmark. Net settlement = only the difference flows. A plain vanilla IRS involves ONE currency and ONE notional.',
        },
        {
            'num': 'QK3', 'type': 'GAP FILL DROPDOWN', 'ear': 'CALCULATE',
            'video_ref': 'SWP.2 Video 2',
            'q': 'SWP.2 Video 2 presents a multi-period IRS example. Complete the settlement logic.',
            'paragraph': 'An IRS has notional Rs 100 crore, fixed rate 7%, semi-annual settlements. Period 1: MIBOR = 6%. The fixed payer pays _1_ crore more than the floating payer receives, so the fixed payer makes a _2_ payment of Rs 0.5 crore. Period 2: MIBOR rises to 8%. Now floating payer owes Rs _3_ crore; fixed payer owes Rs 3.5 crore. Net: floating payer pays Rs _4_ crore to the fixed payer. Period 3: MIBOR = 7%. Net settlement is _5_ because both payments are _6_.',
            'gaps': [
                ('_1_', ['0.5', '1.0', '3.5', '7.0'], '0.5'),
                ('_2_', ['net outflow (pays)', 'net inflow (receives)', 'zero', 'gross'], 'net outflow (pays)'),
                ('_3_', ['4', '3.5', '6', '7'], '4'),
                ('_4_', ['0.5', '3.5', '4', '1.0'], '0.5'),
                ('_5_', ['zero', 'Rs 3.5 crore', 'Rs 7 crore', 'Rs 0.5 crore'], 'zero'),
                ('_6_', ['equal', 'different', 'variable', 'fixed'], 'equal'),
            ],
            'correct_text': '_1_ 0.5 | _2_ net outflow (pays) | _3_ 4 | _4_ 0.5 | _5_ zero | _6_ equal',
            'reasoning': 'From Video 2: Semi-annual payment = rate × 100 Cr × 0.5. Period 1: 7%→Rs3.5Cr, 6%→Rs3Cr, net = fixed payer pays Rs0.5Cr. Period 2: 7%→Rs3.5Cr, 8%→Rs4Cr, net = floating payer pays Rs0.5Cr. Period 3: both equal Rs3.5Cr, net = zero.',
        },
        {
            'num': 'QK4', 'type': 'TABLE SELECTION', 'ear': 'EXPLAIN',
            'video_ref': 'SWP.2 Videos 1–2',
            'q': 'A corporate enters a receive-fixed, pay-floating IRS. Classify each outcome for this corporate.',
            'col_headers': ['Benefits corporate (positive MTM)', 'Hurts corporate (negative MTM)', 'No effect'],
            'rows': [
                ('MIBOR falls to 5% (fixed rate was 7%)', '', 'CORRECT', ''),
                ('MIBOR rises to 9% (fixed rate was 7%)', 'CORRECT', '', ''),
                ('No change in MIBOR', '', '', 'CORRECT'),
                ('MIBOR equals the fixed swap rate exactly', '', '', 'CORRECT'),
                ('RBI unexpectedly cuts rates by 100 bps', '', 'CORRECT', ''),
            ],
            'correct_text': 'MIBOR falls → hurts (pay more floating than receive fixed) | MIBOR rises → benefits (receive fixed > pay floating) | MIBOR = fixed rate → no effect | RBI cuts → rates fall → hurts',
            'reasoning': 'From Video 2: Receive-fixed party benefits when rates FALL (their fixed receipt is above market). Receive-fixed party is hurt when rates RISE (they pay more floating than they receive). This is the OPPOSITE of the pay-fixed party analysis in the video.',
        },
    ],
    'SWP.3': [
        {
            'num': 'QK1', 'type': 'MATCHING PAIRS', 'ear': 'DEFINE',
            'video_ref': 'SWP.3 Video 1',
            'q': 'Match each currency swap concept to its correct description.',
            'col_a': ['Principal Exchange at Inception', 'Principal Exchange at Maturity', 'Periodic Interest', 'IRS vs Currency Swap', 'Use Case for Indian Corporate', 'Exchange Rate Risk on Principal'],
            'col_b': [
                'Party A gives Rs 700 crore; Party B gives $100 million at prevailing spot rate',
                'Original principals returned at the ORIGINAL exchange rate regardless of spot',
                'Each party pays interest in the currency they RECEIVED — their obligation currency',
                'IRS: same currency, no principal exchange | Currency Swap: different currencies, principal exchanged',
                'Issue USD bonds cheaply abroad; swap to INR obligations matching home revenues',
                'Eliminated — principal re-exchanged at original rate locks in currency exposure',
            ],
            'correct_text': 'All 6 pairs match as described — inception principal exchange locks rate; maturity re-exchange at same rate; periodic interest in received currency; currency swap principal moves unlike IRS; Indian corp accesses cheap USD and swaps to INR; FX risk on principal eliminated.',
            'reasoning': 'From Videos 1 and 2: The critical distinction from IRS is that principal IS exchanged at both inception and maturity. Exchange at maturity uses the ORIGINAL rate — eliminating FX risk on principal. Indian corporates use this to access cheaper foreign capital while maintaining INR obligations.',
        },
        {
            'num': 'QK2', 'type': 'TRUE/FALSE MULTI', 'ear': 'EXPLAIN',
            'video_ref': 'SWP.3 Videos 1–2',
            'q': 'Based on SWP.3 Videos 1 and 2, mark each statement True or False.',
            'statements': [
                ('In a currency swap, the principal amounts are exchanged at inception and returned at maturity', True),
                ('The exchange rate at maturity in a currency swap is always the prevailing spot rate on that date', False),
                ('Currency swaps eliminate FX risk on the periodic interest payments', True),
                ('An Indian company can use a currency swap to access cheaper USD bonds while maintaining INR obligations', True),
                ('Currency swaps carry larger counterparty exposure than IRS because principal is actually exchanged', True),
                ('The RBI does not participate in any form of currency swap arrangement', False),
            ],
            'correct_true': [1, 3, 4, 5],
            'correct_false': [2, 6],
            'correct_text': 'True: 1, 3, 4, 5 | False: 2, 6',
            'reasoning': 'Statement 2 is false — maturity exchange uses the ORIGINAL inception rate. Statement 6 is false — the RBI has bilateral currency swap arrangements with the Bank of Japan and others for FX reserve support.',
        },
        {
            'num': 'QK3', 'type': 'MCQ INLINE', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.3 Video 2',
            'q': 'An Indian IT company issues a $100 million bond in the US at 4%, when equivalent INR borrowing would cost 8%. They enter a currency swap: receive $100M, pay INR 7% on Rs 700 crore. What is the PRIMARY reason for this transaction?',
            'options': [
                'A) To speculate that the USD will weaken against the INR over the bond tenor',
                'B) To access cheaper USD capital from international markets while converting the obligation to INR — matching revenues and eliminating currency mismatch on debt service',
                'C) To take advantage of arbitrage between Indian and US equity markets',
                'D) To increase their reported earnings by booking a foreign exchange gain',
            ],
            'correct': 'B',
            'correct_text': 'B) Access cheaper USD capital + convert obligation to INR — eliminate currency mismatch',
            'reasoning': 'From Video 1: "They access cheap dollar capital from international investors, but their actual ongoing obligation is in rupees — matching their INR revenues. No currency mismatch, no FX risk on cash flows." The key objective is cheapest-cost financing with no currency exposure on debt service.',
        },
        {
            'num': 'QK4', 'type': 'GAP FILL DROPDOWN', 'ear': 'DESCRIBE',
            'video_ref': 'SWP.3 Videos 1–2',
            'q': 'Complete the description of currency swap cash flows from SWP.3.',
            'paragraph': 'At _1_, Party A delivers Rs 700 crore to Party B and receives $100 million (rate: 70). Each year, Party A pays _2_ on $100M; Party B pays _3_ on Rs 700 crore. At _4_, Party A returns $100 million and receives back Rs 700 crore — at the _5_ exchange rate of 70, regardless of where spot USD/INR is then. This eliminates _6_ on the principal entirely.',
            'gaps': [
                ('_1_', ['inception', 'maturity', 'the first coupon date', 'settlement'], 'inception'),
                ('_2_', ['4% in USD ($4M)', '8% in INR (Rs 56 Cr)', '7% in INR', '4% in INR'], '4% in USD ($4M)'),
                ('_3_', ['8% in INR (Rs 56 Cr)', '4% in USD', '7% in USD', '8% in USD'], '8% in INR (Rs 56 Cr)'),
                ('_4_', ['maturity', 'inception', 'each coupon date', 'the first year'], 'maturity'),
                ('_5_', ['original', 'current spot', 'forward', 'RBI reference'], 'original'),
                ('_6_', ['FX risk (currency risk)', 'interest rate risk', 'credit risk', 'liquidity risk'], 'FX risk (currency risk)'),
            ],
            'correct_text': '_1_ inception | _2_ 4% USD ($4M) | _3_ 8% INR (Rs 56 Cr) | _4_ maturity | _5_ original | _6_ FX risk (currency risk)',
            'reasoning': 'From Video 2: Principal flows at inception and maturity. Interest paid in the received currency. Maturity re-exchange at original rate regardless of spot rate — this is the mechanism that locks out FX risk on principal.',
        },
    ],
    'SWP.4': [
        {
            'num': 'QK1', 'type': 'TABLE SELECTION', 'ear': 'DEFINE',
            'video_ref': 'SWP.4 Video 1',
            'q': 'Classify each statement as describing the Protection BUYER, Protection SELLER, or BOTH parties in a CDS.',
            'col_headers': ['Protection BUYER only', 'Protection SELLER only', 'BOTH'],
            'rows': [
                ('Pays a regular premium (CDS spread) to the other party', 'CORRECT', '', ''),
                ('Receives the face value of the bond if a credit event occurs', 'CORRECT', '', ''),
                ('Holds a position in the reference entity\'s credit risk', '', '', 'CORRECT'),
                ('Collects the CDS spread payments each quarter', '', 'CORRECT', ''),
                ('Benefits if the reference entity does NOT default', '', 'CORRECT', ''),
                ('Bears risk if the reference entity defaults', '', '', 'CORRECT'),
            ],
            'correct_text': 'Buyer: pays premium, receives payment on default | Seller: collects spread, benefits from no default | Both: exposed to reference entity credit, both bear risk',
            'reasoning': 'From Video 1: Protection buyer PAYS premium, RECEIVES on default. Protection seller RECEIVES premium, PAYS on default. Both parties are exposed to the credit quality of the reference entity in opposite directions.',
        },
        {
            'num': 'QK2', 'type': 'TRUE/FALSE MULTI', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.4 Videos 1–2',
            'q': 'Based on SWP.4 Videos 1 and 2, mark each statement about CDS and credit events True or False.',
            'statements': [
                ('A CDS triggers payment only when a defined credit event occurs — not merely when the reference entity\'s credit rating is downgraded', True),
                ('Bankruptcy, failure to pay, and debt restructuring are recognised ISDA credit events', True),
                ('A "naked" CDS requires the buyer to own the underlying bond before purchasing protection', False),
                ('The CDS spread is quoted in basis points and reflects the market\'s assessment of default probability', True),
                ('After 2008, standardised CDS contracts in major markets must now be centrally cleared', True),
                ('AIG\'s 2008 failure was caused by purchasing too many CDS protection contracts', False),
            ],
            'correct_true': [1, 2, 4, 5],
            'correct_false': [3, 6],
            'correct_text': 'True: 1, 2, 4, 5 | False: 3, 6',
            'reasoning': 'Statement 3 is false — naked CDS means no underlying bond required (that is the definition). Statement 6 is false — AIG SOLD CDS protection; it was overwhelmed when required to pay out on defaults it had guaranteed.',
        },
        {
            'num': 'QK3', 'type': 'MCQ INLINE', 'ear': 'EXPLAIN',
            'video_ref': 'SWP.4 Video 2',
            'q': 'SWP.4 Video 2 describes the role of CDS in the 2008 financial crisis. Which statement BEST explains how CDS amplified systemic risk?',
            'options': [
                'A) Banks used CDS to buy insurance on their own bonds, creating too many protection buyers',
                'B) AIG and others sold massive volumes of CDS protection on CDOs backed by subprime mortgages, without holding adequate capital. When housing prices fell broadly, all CDS triggered simultaneously and AIG could not honour its obligations.',
                'C) The CDS market was too small — insufficient protection was available for bond holders who needed it',
                'D) Regulators accidentally included CDS in the definition of exchange-traded securities, causing panic selling',
            ],
            'correct': 'B',
            'correct_text': 'B) AIG sold $500B in CDS protection on CDOs; mass simultaneous triggering when housing fell; AIG could not pay',
            'reasoning': 'From Video 2: "AIG had written $500 billion in CDS protection it could not honour. The US government was forced to bail out AIG to prevent complete financial collapse." The failure was on the SELLER side — too much protection sold without capital, on correlated underlying assets that all defaulted together.',
        },
        {
            'num': 'QK4', 'type': 'GAP FILL DROPDOWN', 'ear': 'DESCRIBE',
            'video_ref': 'SWP.4 Videos 1–2',
            'q': 'Complete the description of CDS structure and post-2008 reforms from SWP.4.',
            'paragraph': 'A CDS is like _1_ on a bond. The protection buyer pays the CDS _2_ — quoted in basis points — to the seller each quarter. If a credit event occurs, the seller pays the buyer the _3_ of the bond. The credit event is declared by the ISDA _4_. Post-2008, standardised CDS must be cleared through a _5_ to eliminate bilateral counterparty risk. "Naked" CDS — where the buyer has no underlying bond — has been _6_ for sovereign debt in the EU since 2012.',
            'gaps': [
                ('_1_', ['insurance', 'a loan', 'an equity option', 'a futures contract'], 'insurance'),
                ('_2_', ['spread (premium)', 'notional', 'coupon', 'margin'], 'spread (premium)'),
                ('_3_', ['face value', 'market value', 'recovery value', 'book value'], 'face value'),
                ('_4_', ['determinations committee', 'central bank', 'rating agency', 'stock exchange'], 'determinations committee'),
                ('_5_', ['central counterparty (CCP)', 'rating agency', 'central bank', 'stock exchange'], 'central counterparty (CCP)'),
                ('_6_', ['banned', 'encouraged', 'mandated', 'subsidised'], 'banned'),
            ],
            'correct_text': '_1_ insurance | _2_ spread (premium) | _3_ face value | _4_ determinations committee | _5_ CCP | _6_ banned',
            'reasoning': 'From Videos 1 and 2: CDS = insurance analogy. Buyer pays spread (premium). On credit event = pay face value. ISDA determinations committee decides credit events. Post-2008 Dodd-Frank/EMIR: central clearing mandatory. EU banned naked sovereign CDS in 2012.',
        },
    ],
    'SWP.5': [
        {
            'num': 'QK1', 'type': 'MCQ INLINE', 'ear': 'EXPLAIN',
            'video_ref': 'SWP.5 Video 1',
            'q': 'SWP.5 Video 1 explains why a new swap has zero value at inception. Which answer CORRECTLY explains this?',
            'options': [
                'A) Because no money is exchanged at inception, both parties start with equal zero cash positions',
                'B) Because the swap rate is set so that the present value of the fixed leg exactly equals the present value of the expected floating leg — making the trade fair at inception',
                'C) Because swaps are OTC instruments with no market value until the first settlement date',
                'D) Because regulators require all swaps to be issued at par value of zero',
            ],
            'correct': 'B',
            'correct_text': 'B) Swap rate is set so PV(fixed) = PV(floating) — fair exchange at inception means zero net value',
            'reasoning': 'From Video 1: "The swap rate is set so that the present value of the fixed cash flows exactly equals the present value of the expected floating cash flows. It is a fair exchange. Nobody is getting a better deal than the other at inception." Zero value = fair pricing, not absence of value.',
        },
        {
            'num': 'QK2', 'type': 'TRUE/FALSE MULTI', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.5 Videos 1–2',
            'q': 'Based on SWP.5, mark each statement about swap valuation True or False.',
            'statements': [
                ('A swap\'s mark-to-market value is zero at inception but changes as interest rates move', True),
                ('If market rates rise above the fixed rate in a pay-fixed swap, the fixed payer\'s MTM becomes negative', False),
                ('CVA (Credit Valuation Adjustment) adjusts swap value to reflect the probability that your counterparty defaults', True),
                ('To exit a swap before maturity, the only option is to wait for it to expire', False),
                ('The passage of time (fewer remaining cash flows) affects the MTM value of a swap', True),
                ('The OIS (Overnight Index Swap) curve is used in India to derive discount factors for swap valuation', True),
            ],
            'correct_true': [1, 3, 5, 6],
            'correct_false': [2, 4],
            'correct_text': 'True: 1, 3, 5, 6 | False: 2, 4',
            'reasoning': 'Statement 2 is false — if rates rise above the fixed rate, the FIXED PAYER benefits (positive MTM) because they pay below-market rates and receive above-market floating. Statement 4 is false — Video 2 describes three exit options: offsetting swap, early termination, or novation.',
        },
        {
            'num': 'QK3', 'type': 'GAP FILL DROPDOWN', 'ear': 'DESCRIBE',
            'video_ref': 'SWP.5 Video 1',
            'q': 'Complete the explanation of swap MTM drivers from SWP.5 Video 1.',
            'paragraph': 'After a swap is executed, its value changes with four factors. First, _1_ rate changes — a parallel shift up or down in rates is the dominant driver. Second, changes in _2_ shape — steepening or flattening affects long-dated swaps more. Third, _3_ decay — each settlement date that passes reduces the number of remaining cash flows to be discounted. Fourth, the credit quality of the _4_ — if they deteriorate, your positive MTM is worth less on a risk-adjusted basis, captured by the _5_ adjustment (CVA). A swap with _6_ years remaining is more rate-sensitive than one with 1 year remaining.',
            'gaps': [
                ('_1_', ['interest', 'equity', 'commodity', 'property'], 'interest'),
                ('_2_', ['yield curve', 'equity index', 'credit spread', 'FX rate'], 'yield curve'),
                ('_3_', ['time', 'credit', 'equity', 'currency'], 'time'),
                ('_4_', ['counterparty', 'reference entity', 'central bank', 'regulator'], 'counterparty'),
                ('_5_', ['CVA (credit valuation)', 'DVA (debt valuation)', 'PVA (price valuation)', 'LVA (liquidity valuation)'], 'CVA (credit valuation)'),
                ('_6_', ['10', '1', '0.5', '0.25'], '10'),
            ],
            'correct_text': '_1_ interest | _2_ yield curve | _3_ time | _4_ counterparty | _5_ CVA | _6_ 10',
            'reasoning': 'From Video 1: Four MTM drivers: 1) interest rate level 2) yield curve shape 3) time passage 4) counterparty credit quality. CVA adjusts for counterparty default probability. Longer-dated swaps have higher rate sensitivity (more cash flows to discount).',
        },
        {
            'num': 'QK4', 'type': 'MATCHING PAIRS', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.5 Videos 1–2',
            'q': 'Match each swap exit or valuation concept to its correct description.',
            'col_a': ['Swap Rate', 'Mark-to-Market (MTM)', 'CVA', 'Novation', 'Offsetting Swap', 'Early Termination'],
            'col_b': [
                'The fixed rate that makes PV(fixed) = PV(floating) at inception — sets the trade as fair',
                'Daily valuation of the swap based on current market rates and remaining cash flows',
                'Adjustment to MTM reflecting probability of counterparty default before maturity',
                'Transfer of your swap position to a third party with original counterparty\'s consent',
                'Enter a reverse swap — pay floating, receive fixed — to neutralise position without closing original',
                'Agree with counterparty to cancel the swap and pay/receive current MTM as cash settlement',
            ],
            'correct_text': 'All 6 matches as described — swap rate = fair pricing rate | MTM = current valuation | CVA = credit adjustment | Novation = assignment to third party | Offsetting = reverse swap | Early termination = cancel for cash MTM',
            'reasoning': 'From Videos 1 and 2: The swap rate makes the trade fair at inception. MTM fluctuates daily. CVA adjusts for counterparty credit. Three exit routes: offsetting swap, early termination (pay/receive MTM), or novation (transfer to third party with consent).',
        },
    ],
    'SWP.6': [
        {
            'num': 'QK1', 'type': 'MCQ INLINE', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.6 Video 1',
            'q': 'SWP.6 Video 1 explains how banks use IRS to manage balance sheet risk. A bank has Rs 500 crore of 5-year fixed-rate loans funded by 3-month floating deposits. Which IRS position eliminates the bank\'s interest rate mismatch?',
            'options': [
                'A) Pay-fixed IRS: the bank pays fixed and receives floating — adding more fixed-rate exposure',
                'B) Receive-fixed IRS: the bank receives fixed (offsetting fixed-rate loans) and pays floating (matching deposit costs) — neutralising the mismatch',
                'C) The bank should not use swaps — it should instead only issue fixed-rate deposits to match its fixed-rate loans',
                'D) A currency swap that converts the rupee loans into dollar obligations',
            ],
            'correct': 'B',
            'correct_text': 'B) Receive-fixed IRS — receives fixed (matches loan income), pays floating (matches deposit cost)',
            'reasoning': 'From Video 1: "Banks enter receive-fixed IRS to hedge this mismatch. They receive fixed from the swap, offsetting their fixed-rate loans, while the floating payment on the swap matches their floating deposit costs." Pay-fixed would ADD to exposure not reduce it.',
        },
        {
            'num': 'QK2', 'type': 'TRUE/FALSE MULTI', 'ear': 'EXPLAIN',
            'video_ref': 'SWP.6 Videos 1–2',
            'q': 'Based on SWP.6, mark each statement about swap applications and risks True or False.',
            'statements': [
                ('Central clearing through CCIL eliminates bilateral counterparty risk in standardised Indian OTC derivatives', True),
                ('Hedge accounting under IFRS 9 always records swap gains and losses directly in profit and loss', False),
                ('A pay-fixed IRS position is equivalent to a leveraged short bond position for a speculative hedge fund', True),
                ('Counterparty risk is smaller in long-dated currency swaps than in short-dated IRS because more payments remain', False),
                ('Corporate liability management using swaps can convert fixed-rate debt to floating without refinancing the original bonds', True),
                ('Initial margin and variation margin must be posted daily when a swap is centrally cleared', True),
            ],
            'correct_true': [1, 3, 5, 6],
            'correct_false': [2, 4],
            'correct_text': 'True: 1, 3, 5, 6 | False: 2, 4',
            'reasoning': 'Statement 2 is false — under a cash flow hedge (IFRS 9), the effective portion goes to OCI, not P&L directly. Statement 4 is false — longer-dated swaps carry GREATER counterparty risk because more cash flows remain at risk. CCIL clears standardised OTC derivatives in India.',
        },
        {
            'num': 'QK3', 'type': 'GAP FILL DROPDOWN', 'ear': 'DESCRIBE',
            'video_ref': 'SWP.6 Video 2',
            'q': 'Complete the description of central clearing and its role in swap risk management from SWP.6 Video 2.',
            'paragraph': 'In central clearing, the _1_ becomes the buyer to every seller and the seller to every buyer — eliminating _2_ between the original swap counterparties. In India, _3_ provides central clearing for IRS and OIS. To participate, each party must post _4_ margin (upfront) and _5_ margin (daily, reflecting MTM changes). This regime was mandated post-2008 to prevent the type of systemic failure seen with _6_, which had written unhedged CDS protection it could not honour.',
            'gaps': [
                ('_1_', ['Central Counterparty (CCP)', 'SEBI', 'Reserve Bank of India', 'ISDA'], 'Central Counterparty (CCP)'),
                ('_2_', ['bilateral counterparty risk', 'market risk', 'liquidity risk', 'currency risk'], 'bilateral counterparty risk'),
                ('_3_', ['CCIL (Clearing Corporation of India)', 'NSE Clearing', 'BSE', 'RBI directly'], 'CCIL (Clearing Corporation of India)'),
                ('_4_', ['initial', 'variation', 'maintenance', 'performance'], 'initial'),
                ('_5_', ['variation', 'initial', 'regulatory', 'credit'], 'variation'),
                ('_6_', ['AIG', 'Lehman Brothers', 'Bear Stearns', 'Goldman Sachs'], 'AIG'),
            ],
            'correct_text': '_1_ CCP | _2_ bilateral counterparty risk | _3_ CCIL | _4_ initial | _5_ variation | _6_ AIG',
            'reasoning': 'From Video 2: CCP interposes itself as counterparty to all. CCIL is the Indian clearing corp. Two types of margin: initial (upfront, covers potential future exposure) and variation (daily, covers MTM changes). AIG is the canonical example of unhedged CDS protection failure.',
        },
        {
            'num': 'QK4', 'type': 'MATCHING PAIRS', 'ear': 'IDENTIFY',
            'video_ref': 'SWP.6 Videos 1–2',
            'q': 'Match each swap application to its primary purpose and user type.',
            'col_a': ['Bank using receive-fixed IRS', 'Corporate entering pay-fixed IRS on floating debt', 'Macro hedge fund paying fixed in OIS', 'Corporate doing currency swap on USD bond', 'CDS protection buyer (bank)', 'Corporate converting fixed bond to floating via IRS'],
            'col_b': [
                'Hedge fixed-rate loan / floating deposit mismatch on balance sheet',
                'Convert floating rate loan to fixed — lock in predictable interest cost',
                'Speculate that RBI will raise rates — pay below-market fixed, receive rising floating',
                'Access cheap foreign capital while maintaining domestic currency obligations',
                'Reduce credit exposure to a borrower without selling the loan',
                'Liability management — benefit from expected rate decline without refinancing',
            ],
            'correct_text': 'All 6 pairs match as described — bank ALM | corporate hedging | macro speculation | currency swap access | CDS credit hedge | liability management',
            'reasoning': 'From Videos 1 and 2: Banks use receive-fixed to fix NIM. Corporates use pay-fixed to hedge floating debt. Macro funds use pay-fixed to speculate on rate rises. Currency swaps for cross-border access. CDS to hedge credit exposure without loan sales. Liability management to adjust duration.',
        },
    ],
}

# ─────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────
def build_assessment():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DERIVATIVES: SWAPS — Knowledge Assessment')
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('CIBOP Module 7 | Version 2 | 24 Questions across 6 UORs')
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)

    doc.add_paragraph()

    # Purpose table
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'
    meta = [('Purpose', 'Test conceptual understanding of SWP.1–SWP.6 | Assess interview readiness across all UOR EARs'),
            ('Question Types', 'MCQ Inline | True/False Multi | Gap Fill Dropdown | Matching Pairs | Table Selection')]
    for r, (k, v) in enumerate(meta):
        cell0 = t.rows[r].cells[0]
        set_cell_bg(cell0, '1E3A5F')
        set_col_width(cell0, 1.5)
        p2 = cell0.paragraphs[0]
        run2 = p2.add_run(k)
        run2.bold = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(255,255,255)
        cell1 = t.rows[r].cells[1]
        set_cell_bg(cell1, 'EEF4FB')
        set_col_width(cell1, 6.0)
        p3 = cell1.paragraphs[0]
        run3 = p3.add_run(v)
        run3.font.size = Pt(9)

    doc.add_paragraph()

    UOR_TITLES = {
        'SWP.1': 'What is a Swap?',
        'SWP.2': 'Interest Rate Swaps (IRS)',
        'SWP.3': 'Currency Swaps',
        'SWP.4': 'Credit Default Swaps (CDS)',
        'SWP.5': 'Swap Pricing & Valuation Basics',
        'SWP.6': 'Swap Applications & Risk Management',
    }

    for uor_id, questions in QUESTIONS.items():
        doc.add_page_break()
        # UOR section header
        p_hdr = doc.add_paragraph()
        run_hdr = p_hdr.add_run(f'{uor_id} — {UOR_TITLES[uor_id]} (Knowledge Questions)')
        run_hdr.bold = True
        run_hdr.font.size = Pt(13)
        run_hdr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        for q in questions:
            doc.add_paragraph()
            # Q header row
            add_q_header(doc, q['num'], q['type'], q['ear'], uor_id, q['video_ref'])

            # Question text
            p_q = doc.add_paragraph()
            p_q.paragraph_format.left_indent = Inches(0.2)
            run_q = p_q.add_run('Question: ')
            run_q.bold = True
            run_q.font.size = Pt(10)
            run_q.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            run_q2 = p_q.add_run(q['q'])
            run_q2.font.size = Pt(10)

            qtype = q['type']

            if qtype == 'MCQ INLINE':
                for opt in q['options']:
                    p_opt = doc.add_paragraph()
                    p_opt.paragraph_format.left_indent = Inches(0.4)
                    run_opt = p_opt.add_run(opt)
                    run_opt.font.size = Pt(9.5)
                    if opt.startswith(f'{q["correct"]})'):
                        run_opt.bold = True
                        run_opt.font.color.rgb = RGBColor(0x00, 0x7A, 0x00)
                answer_box(doc, q['correct_text'], q['reasoning'])

            elif qtype == 'TRUE/FALSE MULTI':
                tf_table = doc.add_table(rows=len(q['statements'])+1, cols=3)
                tf_table.style = 'Table Grid'
                hdrs = ['Statement', 'True', 'False']
                widths_tf = [5.0, 0.8, 0.8]
                for i, (h, w) in enumerate(zip(hdrs, widths_tf)):
                    cell = tf_table.rows[0].cells[i]
                    set_cell_bg(cell, '1E3A5F')
                    set_col_width(cell, w)
                    p2 = cell.paragraphs[0]
                    run2 = p2.add_run(h)
                    run2.bold = True
                    run2.font.size = Pt(9)
                    run2.font.color.rgb = RGBColor(255,255,255)
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_i, (stmt, is_true) in enumerate(q['statements']):
                    row = tf_table.rows[r_i+1]
                    fill = 'FFFFFF' if r_i % 2 == 0 else 'F5F5F5'
                    for i, w in enumerate(widths_tf):
                        set_col_width(row.cells[i], w)
                        set_cell_bg(row.cells[i], fill)
                    row.cells[0].paragraphs[0].add_run(stmt).font.size = Pt(9)
                    mark = 'CORRECT' if is_true else ''
                    row.cells[1].paragraphs[0].add_run(mark).font.size = Pt(9)
                    mark2 = 'CORRECT' if not is_true else ''
                    row.cells[2].paragraphs[0].add_run(mark2).font.size = Pt(9)
                doc.add_paragraph()
                answer_box(doc, q['correct_text'], q['reasoning'])

            elif qtype == 'GAP FILL DROPDOWN':
                p_para = doc.add_paragraph()
                p_para.paragraph_format.left_indent = Inches(0.2)
                run_para = p_para.add_run(q['paragraph'])
                run_para.font.size = Pt(9.5)
                run_para.italic = True

                gap_table = doc.add_table(rows=len(q['gaps'])+1, cols=3)
                gap_table.style = 'Table Grid'
                for i, (h, w) in enumerate(zip(['Gap', 'Options (select one)', 'Correct Answer'], [0.8, 4.0, 2.5])):
                    cell = gap_table.rows[0].cells[i]
                    set_cell_bg(cell, '005F5F')
                    set_col_width(cell, w)
                    p2 = cell.paragraphs[0]
                    run2 = p2.add_run(h)
                    run2.bold = True
                    run2.font.size = Pt(9)
                    run2.font.color.rgb = RGBColor(255,255,255)
                for r_i, (gap_id, options, answer) in enumerate(q['gaps']):
                    row = gap_table.rows[r_i+1]
                    fill = 'FFFFFF' if r_i % 2 == 0 else 'F5F5F5'
                    for col_i, (val, w) in enumerate(zip([gap_id, ' / '.join(options), f'✓ {answer}'], [0.8, 4.0, 2.5])):
                        set_col_width(row.cells[col_i], w)
                        set_cell_bg(row.cells[col_i], fill)
                        run = row.cells[col_i].paragraphs[0].add_run(val)
                        run.font.size = Pt(9)
                        if col_i == 2:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x00, 0x7A, 0x00)
                doc.add_paragraph()
                answer_box(doc, q['correct_text'], q['reasoning'])

            elif qtype == 'MATCHING PAIRS':
                two_col_table(doc, 'Column A — Item', q['col_a'], 'Column B — Match', q['col_b'])
                answer_box(doc, q['correct_text'], q['reasoning'])

            elif qtype == 'TABLE SELECTION':
                sel_table = doc.add_table(rows=len(q['rows'])+1, cols=len(q['col_headers'])+1)
                sel_table.style = 'Table Grid'
                # Header
                set_cell_bg(sel_table.rows[0].cells[0], '1E3A5F')
                set_col_width(sel_table.rows[0].cells[0], 3.0)
                sel_table.rows[0].cells[0].paragraphs[0].add_run('Scenario').font.color.rgb = RGBColor(255,255,255)
                for i, col_hdr in enumerate(q['col_headers']):
                    cell = sel_table.rows[0].cells[i+1]
                    set_cell_bg(cell, '1E3A5F')
                    set_col_width(cell, 5.3 / len(q['col_headers']))
                    run = cell.paragraphs[0].add_run(col_hdr)
                    run.bold = True
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(255,255,255)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_i, row_data in enumerate(q['rows']):
                    row = sel_table.rows[r_i+1]
                    fill = 'FFFFFF' if r_i % 2 == 0 else 'F5F5F5'
                    set_cell_bg(row.cells[0], fill)
                    set_col_width(row.cells[0], 3.0)
                    row.cells[0].paragraphs[0].add_run(row_data[0]).font.size = Pt(9)
                    for i in range(len(q['col_headers'])):
                        cell = row.cells[i+1]
                        set_cell_bg(cell, fill)
                        set_col_width(cell, 5.3 / len(q['col_headers']))
                        val = row_data[i+1]
                        run = cell.paragraphs[0].add_run(val)
                        run.font.size = Pt(9)
                        if val == 'CORRECT':
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x00, 0x7A, 0x00)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                answer_box(doc, q['correct_text'], q['reasoning'])

    # Footer
    for sec in doc.sections:
        footer = sec.footer
        p = footer.paragraphs[0]
        p.clear()
        run = p.add_run('CIBOP PPT 07 — Derivatives: SWAPS  |  Knowledge Assessment  |  Confidential — Internal Use Only')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')
    import os
    print(f'File size: {os.path.getsize(OUTPUT_PATH):,} bytes')

if __name__ == '__main__':
    build_assessment()
    print('Knowledge Assessment DONE.')
