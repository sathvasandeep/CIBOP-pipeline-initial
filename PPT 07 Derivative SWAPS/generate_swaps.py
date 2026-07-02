#!/usr/bin/env python3
"""Generate SWAPS training materials for CIBOP Module 7"""
import subprocess, sys

# Install required packages
subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', 'python-pptx', '-q'], check=True)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_DIR = '/Users/sandeeppr/CIBOP/PPT 07 Derivative SWAPS/'

# ─────────────────────────────────────────────
# UOR DATA
# ─────────────────────────────────────────────
UORS = [
    {
        'id': 'SWP.1',
        'title': 'What is a Swap?',
        'objective': 'Define a swap and explain how two parties exchange cash flows based on different terms.',
        'competency': 'Swap Fundamentals',
        'subcompetencies': [
            'SC1.1 Define a swap as an OTC derivative',
            'SC1.2 Identify the two parties in a swap (fixed payer, floating payer)',
            'SC1.3 Explain why swaps are used (hedging, speculation, comparative advantage)',
        ],
        'ears': 'Define · Identify · Explain',
        'timeline': '10 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Define & Identify',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at desk; split-screen showing two firms exchanging arrows labelled "Fixed" and "Floating"',
                        'onscreen': 'What is a Swap?\nTwo parties. Two obligations. One agreement.',
                        'vo': 'Imagine you have a loan at a floating interest rate. Your competitor has a loan at a fixed rate. You both want what the other has. A swap is simply an agreement to exchange those cash flow obligations — no loan transfer, no asset movement. Just a contractual exchange of payment streams.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; diagram: "Party A (Fixed Payer)" → "Party B (Floating Payer)" with arrows both ways',
                        'onscreen': 'Swap Structure\nParty A pays Fixed → Party B\nParty B pays Floating → Party A',
                        'vo': 'In every swap there are exactly two counterparties. One pays a fixed rate — say 7% per annum. The other pays a floating rate — say MIBOR plus a spread. They swap these payment obligations on an agreed notional principal. Neither party physically hands over the principal itself — only the net interest difference changes hands.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'Animation: OTC label, two banks shaking hands, ISDA document overlay',
                        'onscreen': 'Swaps are OTC Derivatives\n• Bilateral contract\n• Governed by ISDA Master Agreement\n• Not exchange-traded',
                        'vo': 'Swaps are over-the-counter derivatives — bilateral contracts negotiated directly between two parties, typically governed by an ISDA Master Agreement. Because they are OTC, the terms are customisable: notional size, tenor, payment frequency, reference rate. This flexibility is both a feature and a risk — and we will get to counterparty risk shortly.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA to camera; three icons: Shield (hedge), Chart (speculate), Scale (comparative advantage)',
                        'onscreen': 'Why Use Swaps?\n1. Hedging\n2. Speculation\n3. Comparative Advantage',
                        'vo': 'Swaps are used for three main reasons. First, hedging — a corporate borrowing at floating rates can swap into fixed to remove interest rate uncertainty. Second, speculation — a trader who believes rates will rise can pay fixed and receive floating to profit from that view. Third, comparative advantage — two firms with different borrowing costs in different markets can each borrow where they have an advantage and swap to achieve their desired exposure at a lower total cost.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Explain (Deeper)',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO at desk; graphic of two companies — "Infra Co" and "Tech Co" — with their borrowing cost tables',
                        'onscreen': 'Comparative Advantage in Swaps\nEach borrows where they have the edge\nThen swap obligations',
                        'vo': 'Here is the classic comparative advantage story. Infra Co can borrow fixed at 8% or floating at MIBOR + 1%. Tech Co can borrow fixed at 10% or floating at MIBOR + 2%. Infra Co has an advantage in fixed markets; Tech Co in floating. Each borrows where they are cheapest, then swaps. Both end up with their preferred rate at a lower all-in cost. That is not magic — that is arbitrage of credit spreads.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA gesturing; timeline arrow showing inception → periodic settlements → maturity',
                        'onscreen': 'Life of a Swap\nInception → Periodic Cash Flows → Maturity\nNo principal exchange (in IRS)',
                        'vo': 'Once a swap is executed, the parties do not transfer money upfront. At each settlement date — typically every three or six months — they calculate what each owes and net it out. Only the difference is paid. If Party A owes Rs 60 lakh in fixed interest and Party B owes Rs 55 lakh in floating interest, Party A simply pays Rs 5 lakh to Party B. Clean, efficient, and bilateral.',
                    },
                    {
                        'num': 7, 'char': 'RYO',
                        'visual': 'Comparison table on screen: Swap vs Forward vs Option vs Futures',
                        'onscreen': 'Swap vs Other Derivatives\nSwap = series of forward contracts\nLonger tenors (1–10 years typical)',
                        'vo': 'One useful mental model: an interest rate swap is economically equivalent to a series of forward rate agreements — one for each settlement date. Each settlement is like a forward settling. The key distinction from futures is that swaps are OTC, customisable, and typically longer-dated — two to ten years is common. Futures are standardised, exchange-traded, and short-dated.',
                    },
                    {
                        'num': 8, 'char': 'ARIA',
                        'visual': 'ARIA smiling; summary slide with three bullets and "Next: Interest Rate Swaps" teaser',
                        'onscreen': 'Summary — SWP.1\n• Swap = exchange of cash flow obligations\n• Two parties: fixed payer and floating payer\n• Used for hedging, speculation, comparative advantage',
                        'vo': 'So: a swap is an OTC agreement to exchange cash flows based on different terms, on a notional principal, over an agreed period. The parties are the fixed payer and the floating payer. The motivations are hedging rate risk, expressing a market view, or exploiting comparative credit advantages. In the next module, we go deep on the most common swap — the plain vanilla interest rate swap.',
                    },
                ],
            },
        ],
        'next_uor': 'Interest Rate Swaps (IRS)',
    },
    {
        'id': 'SWP.2',
        'title': 'Interest Rate Swaps (IRS)',
        'objective': 'Explain the mechanics of a plain vanilla interest rate swap and calculate net cash flows.',
        'competency': 'Interest Rate Swap Mechanics',
        'subcompetencies': [
            'SC2.1 Describe the plain vanilla IRS structure (fixed vs floating)',
            'SC2.2 Identify LIBOR/MIBOR as the floating rate benchmark',
            'SC2.3 Calculate net settlement payments in an IRS',
            'SC2.4 Explain the role of notional principal',
        ],
        'ears': 'Describe · Identify · Calculate · Explain',
        'timeline': '12 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Describe & Identify',
                'scenes': [
                    {
                        'num': 1, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; large diagram: "Fixed Rate Payer" ↔ "Floating Rate Payer" with MIBOR label',
                        'onscreen': 'Plain Vanilla IRS\nFixed Rate Payer ↔ Floating Rate Payer\nNotional: Rs 100 Crore | Tenor: 3 Years',
                        'vo': 'The plain vanilla interest rate swap is the simplest and most traded swap in the world. One party pays a fixed rate — agreed at the start and unchanged throughout. The other pays a floating rate — typically MIBOR in India or LIBOR historically in global markets — which resets each period. The notional principal — say Rs 100 crore — never changes hands. It is purely a reference amount for calculating interest payments.',
                    },
                    {
                        'num': 2, 'char': 'RYO',
                        'visual': 'RYO at screen; MIBOR chart showing rate movements over 3 years',
                        'onscreen': 'MIBOR — The Floating Benchmark\nMumbai Interbank Offer Rate\nResets every 3 or 6 months',
                        'vo': 'MIBOR — the Mumbai Interbank Offer Rate — is the Indian equivalent of LIBOR. It represents the rate at which Indian banks lend to each other in the short-term money market. In an IRS, the floating rate is typically MIBOR flat or MIBOR plus a spread, and it resets every three or six months based on the prevailing market rate. As MIBOR moves with RBI policy rates and market liquidity, the floating leg cash flows vary each period.',
                    },
                    {
                        'num': 3, 'char': 'ARIA',
                        'visual': 'Animation: calculation table — columns: Period, Fixed Payment, Floating Payment (MIBOR), Net Settlement',
                        'onscreen': 'Calculating Net Settlement\nFixed = 7% × Rs 100 Cr × 0.5 = Rs 3.5 Cr\nFloating = MIBOR × Rs 100 Cr × 0.5\nNet = Difference (paid by one party to other)',
                        'vo': 'Let us walk through the arithmetic. Notional: Rs 100 crore. Fixed rate: 7% per annum. Semi-annual settlement. Each period, the fixed payer owes: 7% × 100 crore × 0.5 = Rs 3.5 crore. If MIBOR that period is 6%, the floating payer owes: 6% × 100 crore × 0.5 = Rs 3 crore. Net settlement: fixed payer pays Rs 0.5 crore to floating payer. Only the net flows — clean and efficient.',
                    },
                    {
                        'num': 4, 'char': 'RYO',
                        'visual': 'RYO gesturing at notional principal graphic — large "Rs 100 Cr" with a cross through it',
                        'onscreen': 'Notional Principal\n• Reference amount only\n• NEVER exchanged between parties\n• Used ONLY to calculate interest payments',
                        'vo': 'The notional principal is one of the most misunderstood concepts in swaps. It is not real money that moves. It never changes hands. It is purely a reference number — the base on which interest is calculated. A Rs 100 crore notional swap does not involve Rs 100 crore of cash flow. It involves only the net interest differential, which is typically a tiny fraction of the notional. This is why swaps can represent enormous face values with relatively small cash requirements.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Calculate & Explain',
                'scenes': [
                    {
                        'num': 5, 'char': 'ARIA',
                        'visual': 'ARIA at desk; full worked example table across multiple periods with varying MIBOR',
                        'onscreen': 'Multi-Period IRS Example\nPeriod 1: MIBOR=6% → Net: Fixed pays Rs 0.5 Cr\nPeriod 2: MIBOR=8% → Net: Floating pays Rs 0.5 Cr\nPeriod 3: MIBOR=7% → Net Settlement = Zero',
                        'vo': 'Watch what happens when MIBOR moves. Period 1: MIBOR is 6% — fixed payer pays 7%, floating payer pays 6%, so the fixed payer nets out Rs 0.5 crore. Period 2: MIBOR rises to 8% — now the floating payer pays more, so the floating payer nets Rs 0.5 crore to the fixed payer. Period 3: MIBOR equals the fixed rate exactly — net settlement is zero. The swap is effectively a series of bets on where MIBOR will land relative to 7%.',
                    },
                    {
                        'num': 6, 'char': 'RYO',
                        'visual': 'RYO with two-column comparison: "Corporate with floating loan" vs "After entering IRS"',
                        'onscreen': 'Why Corporates Use IRS\nConvert floating-rate debt → effective fixed rate\nLock in certainty for budgeting and planning',
                        'vo': 'A manufacturing company has borrowed Rs 100 crore from a bank at MIBOR plus 1%. Their CFO is nervous about rising rates. So they enter a swap: pay fixed 7%, receive MIBOR. Now their effective cost is: bank loan cost (MIBOR + 1%) minus MIBOR received from swap, plus fixed 7% paid = 8% fixed total. They have converted their floating obligation into a fixed obligation — at a predictable cost — without renegotiating the original loan.',
                    },
                    {
                        'num': 7, 'char': 'ARIA',
                        'visual': 'ARIA pointing at risk profile graphic — "If MIBOR rises: Fixed payer wins. If MIBOR falls: Floating payer wins."',
                        'onscreen': 'IRS — Who Benefits When?\nRates rise → Fixed payer benefits (pays less relatively)\nRates fall → Floating payer benefits (pays less)',
                        'vo': 'Understanding who wins and who loses in an IRS is critical. If you pay fixed and receive floating, you want rates to go UP — because as MIBOR rises above your fixed rate, you receive more than you pay. Conversely, if you pay floating and receive fixed, you want rates to go DOWN. A speculator who believes RBI will hike rates would enter as a fixed payer. A hedger who has floating debt would enter as a fixed payer to cap their cost.',
                    },
                    {
                        'num': 8, 'char': 'RYO',
                        'visual': 'RYO smiling; summary table with three key takeaways; "Next: Currency Swaps" banner',
                        'onscreen': 'IRS — Key Takeaways\n• Fixed vs floating exchange on notional\n• Net settlement only — no notional transfer\n• Most liquid derivative in Indian OTC markets',
                        'vo': 'To summarise: the plain vanilla IRS involves exchanging fixed and floating interest payments on a notional principal. Only net cash flows are exchanged at each settlement date. MIBOR is the Indian floating benchmark. Corporates use IRS to hedge interest rate exposure; banks and traders use them to express rate views or manage their own balance sheet risk. The IRS market in India, operated through the OIS and MIFOR markets, is among the most liquid OTC derivative markets in the country.',
                    },
                ],
            },
        ],
        'next_uor': 'Currency Swaps',
    },
    {
        'id': 'SWP.3',
        'title': 'Currency Swaps',
        'objective': 'Explain how currency swaps work and identify their use in managing exchange rate risk.',
        'competency': 'Currency Swap Understanding',
        'subcompetencies': [
            'SC3.1 Define a currency swap and distinguish it from an IRS',
            'SC3.2 Explain the exchange of principal at inception and maturity',
            'SC3.3 Identify how corporates use currency swaps to access foreign markets',
            'SC3.4 Describe the cash flow structure of a currency swap',
        ],
        'ears': 'Define · Explain · Identify · Describe',
        'timeline': '11 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Define & Explain',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at screen; world map with INR↔USD arrows between India and USA',
                        'onscreen': 'Currency Swap\nExchange cash flows in DIFFERENT currencies\nNotional principal IS exchanged — unlike IRS',
                        'vo': 'A currency swap takes the IRS concept and adds a currency dimension. Two parties exchange interest payments — and principal — in two different currencies. An Indian company needing dollars and an American company needing rupees might arrange a currency swap: each borrows in their home currency and swaps obligations. The critical difference from an IRS: in a currency swap, the principal amounts are actually exchanged at the start and returned at maturity.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; three-phase diagram: Inception, Periodic, Maturity with cash flow arrows',
                        'onscreen': 'Currency Swap — Three Phases\nInception: Principal exchanged\nPeriodic: Interest payments in respective currencies\nMaturity: Principal re-exchanged at original rate',
                        'vo': 'The structure has three phases. At inception: Party A gives Rs 700 crore to Party B; Party B gives $100 million to Party A — at the prevailing USD/INR rate of 70. During the tenor: Party A pays dollar interest on $100 million; Party B pays rupee interest on Rs 700 crore. At maturity: the original principals are returned at the same original exchange rate — regardless of where spot USD/INR is then. This eliminates currency risk on the principal entirely.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'Side-by-side comparison table: IRS vs Currency Swap',
                        'onscreen': 'IRS vs Currency Swap\nIRS: Same currency, no principal exchange\nCurrency Swap: Different currencies, principal exchanged\nCurrency Swap: Eliminates both FX and interest rate risk',
                        'vo': 'Here is how they differ. An IRS: same currency, same notional — only interest streams swap, principal stays put. A currency swap: two different currencies, and the principal actually exchanges hands at both ends of the transaction. An IRS hedges interest rate risk only. A currency swap hedges both interest rate risk AND currency exchange rate risk simultaneously. For a CFO managing cross-border debt, that is a very powerful combination.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA at desk; case study: "Infosys issues dollar bonds but wants rupee obligations"',
                        'onscreen': 'Corporate Use Case\nIndian IT firm issues USD bond abroad\nEnters currency swap to convert to INR obligations\nResult: Access to cheaper foreign capital + INR cash flow certainty',
                        'vo': 'Consider an Indian technology company that can issue dollar bonds in the US market at 4% — cheaper than domestic rupee borrowing at 8%. They issue the dollar bond, collect $100 million, then enter a currency swap to receive dollars and pay rupees. Net result: they access cheap dollar capital from international investors, but their actual ongoing obligation is in rupees — matching their INR revenues. No currency mismatch, no FX risk on cash flows.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Identify & Describe',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; full cash flow table: USD/INR swap between Indian and US company',
                        'onscreen': 'Full Cash Flow Example\nNotional: $100M / Rs 700 Cr (at 70 USD/INR)\nAnnual: Indian co pays 4% USD = $4M\nUS co pays 8% INR = Rs 56 Cr',
                        'vo': 'Let us build the full picture. Indian Co borrows in the US at 4% and swaps: it receives $100 million, pays rupees at 8%. US Co does the reverse. Each year: Indian Co pays $4 million to US Co; US Co pays Rs 56 crore to Indian Co. At maturity, they re-exchange: Indian Co returns $100 million, US Co returns Rs 700 crore — at that original rate of 70, regardless of whether the spot rate is now 80 or 65. The original rate locks the principal exchange permanently.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA smiling; map of ASEAN companies accessing Japanese yen markets, World Bank currency swap diagram',
                        'onscreen': 'World Bank & Sovereign Currency Swaps\nCentral banks use currency swaps for FX reserves\nEmerging market sovereigns access hard currency',
                        'vo': 'Currency swaps are not just a corporate tool. Central banks — including the RBI — use bilateral currency swap arrangements with other central banks to shore up foreign exchange reserves during stress periods. The RBI has swap lines with the Bank of Japan and other central banks. These provide emergency liquidity in foreign currencies without needing to sell reserves. The mechanics are identical — just at sovereign scale and with political undertones.',
                    },
                    {
                        'num': 7, 'char': 'RYO',
                        'visual': 'RYO at screen; risk matrix: FX risk reduced, Counterparty risk remains, Settlement risk highlighted',
                        'onscreen': 'Risks in Currency Swaps\n• Counterparty default risk (large given principal exchange)\n• Translation risk remains if unhedged residuals\n• Settlement risk — two currencies, two settlement systems',
                        'vo': 'Because principal is actually exchanged in a currency swap, the counterparty exposure is far larger than in an IRS. If your counterparty defaults after you have handed over Rs 700 crore but before you receive your $100 million back at maturity, the loss is enormous. This is why currency swaps between banks are governed by netting agreements and collateral requirements under ISDA/CSA frameworks. The FX risk is eliminated, but the credit risk is amplified.',
                    },
                    {
                        'num': 8, 'char': 'ARIA',
                        'visual': 'ARIA at desk; summary card; "Next: Credit Default Swaps" teaser animation',
                        'onscreen': 'Currency Swap — Summary\n• Principal exchanged at inception and maturity\n• Periodic interest in respective currencies\n• Used to access foreign capital + eliminate FX risk on debt',
                        'vo': 'Currency swaps are elegant structures for managing cross-border financing. Principal exchanges at inception and maturity eliminate the currency risk on the loan itself. Periodic interest payments match the revenue currency of the borrower. Corporates use them to access cheaper foreign funding without taking on currency mismatch. Up next: we move from interest rate and currency risk to credit risk — with the Credit Default Swap.',
                    },
                ],
            },
        ],
        'next_uor': 'Credit Default Swaps (CDS)',
    },
    {
        'id': 'SWP.4',
        'title': 'Credit Default Swaps (CDS)',
        'objective': 'Explain what a CDS is, identify its components, and describe how it transfers credit risk.',
        'competency': 'Credit Derivative Awareness',
        'subcompetencies': [
            'SC4.1 Define a CDS and the protection buyer/seller roles',
            'SC4.2 Identify the credit event that triggers CDS payment',
            'SC4.3 Explain the premium (spread) paid by the protection buyer',
            'SC4.4 Describe how CDS was used/misused during the 2008 financial crisis',
        ],
        'ears': 'Define · Identify · Explain · Describe',
        'timeline': '12 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Define & Identify',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at desk; insurance metaphor graphic — bond investor holding umbrella labelled "CDS"',
                        'onscreen': 'Credit Default Swap\n"Insurance" on a bond or loan\nProtection Buyer pays premium\nProtection Seller pays on default',
                        'vo': 'Think of a CDS as insurance on a bond. You hold a bond issued by a company. You are worried the company might default. You buy a CDS — a credit default swap. You pay a regular premium to the protection seller. If the company defaults — fails to pay, goes bankrupt, or restructures debt — the protection seller pays you the face value of the bond. You have transferred the credit risk without selling the bond itself.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; two-column diagram: Protection Buyer and Protection Seller with premium and default payment arrows',
                        'onscreen': 'CDS Structure\nProtection Buyer → pays CDS spread (e.g. 150 bps p.a.)\nProtection Seller → pays face value if credit event occurs\nReference Entity: the company/bond being insured',
                        'vo': 'The CDS has three key components. One: the reference entity — the company or sovereign whose credit risk is being transferred. Two: the protection buyer — typically a bondholder or bank that wants to reduce credit exposure. Three: the protection seller — typically an insurer, bank, or hedge fund willing to take on that credit risk for a premium. The premium — called the CDS spread — is quoted in basis points per annum on the notional and paid quarterly.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'RYO at screen; list of credit events with examples',
                        'onscreen': 'Credit Events That Trigger CDS\n1. Bankruptcy / Insolvency\n2. Failure to Pay (missed payment)\n3. Debt Restructuring\n4. Repudiation / Moratorium (sovereigns)\n5. Obligation Acceleration',
                        'vo': 'The CDS only pays out on a defined credit event. ISDA has standardised what qualifies: bankruptcy or insolvency, failure to pay a scheduled interest or principal, debt restructuring where terms are materially worsened, or for sovereigns, repudiation or moratorium. When a credit event occurs, the CDS is settled — either physically, where the buyer delivers the defaulted bond and receives face value, or in cash, where the difference between face value and recovery value is paid.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA smiling; CDS spread chart — rising spread = rising perceived default risk',
                        'onscreen': 'CDS Spread = Market\'s View of Default Risk\nHigher spread → market expects higher default probability\nYes Bank CDS (2020): spread spiked before collapse\nAA-rated: 20–50 bps | Distressed: 500–2000 bps',
                        'vo': 'The CDS spread is a real-time market signal of creditworthiness. A tightly rated sovereign might trade at 20 basis points per annum — almost no default risk priced in. A stressed corporate might trade at 500 basis points or more. When Yes Bank faced a crisis in early 2020, its CDS spread blew out dramatically before the RBI intervention. Traders who had bought CDS protection were sitting on enormous profits. It is, quite literally, the market putting a price on default.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Explain & Describe (2008 crisis)',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO serious expression; 2008 crisis timeline graphic — CDO/CDS chain',
                        'onscreen': 'CDS and the 2008 Financial Crisis\nAIG sold $500 billion in CDS protection\nOn mortgage-backed securities\nWhen housing defaulted — AIG could not pay',
                        'vo': 'Here is where CDS went from elegant hedging tool to systemic catastrophe. In the mid-2000s, banks packed subprime mortgages into CDOs — collateralised debt obligations — and AIG and others sold CDS protection on those CDOs in massive volumes. The implicit assumption was that US housing prices would never fall broadly. When they did, every CDS triggered simultaneously. AIG had written $500 billion in CDS protection it could not honour. The US government was forced to bail out AIG to prevent a complete financial collapse.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA at desk; graphic showing "naked CDS" — buyer has no underlying bond exposure',
                        'onscreen': '"Naked" CDS\nBuyer has NO underlying bond\nPure speculation on default\nBanned for sovereign debt in EU post-2012\nLike insuring a house you don\'t own',
                        'vo': 'The 2008 crisis revealed a dark feature of CDS: anyone could buy protection on a company without owning any of its bonds. This is the "naked" CDS — pure speculation on default. Imagine buying fire insurance on your neighbour\'s house. You now have an incentive for it to burn down. Critics argued naked CDS accelerated the collapse of Bear Stearns and Lehman Brothers by allowing speculative short-selling of credit. The EU banned naked sovereign CDS in 2012 for exactly this reason.',
                    },
                    {
                        'num': 7, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; post-2008 CDS reform graphic — Central Clearing, Dodd-Frank, ISDA determinations committee',
                        'onscreen': 'Post-2008 CDS Reforms\n• Dodd-Frank Act: standardised CDS must clear centrally\n• ISDA Determinations Committee: declares credit events\n• Increased transparency via trade repositories',
                        'vo': 'After 2008, regulators overhauled the CDS market. The Dodd-Frank Act in the US and EMIR in Europe mandated that standardised CDS contracts be cleared through central counterparties — eliminating the bilateral counterparty risk that destroyed AIG. ISDA formed determinations committees to make binding decisions on whether credit events have occurred, removing ambiguity. The CDS market is now smaller and more transparent — but it remains a critical tool for credit risk management in banking.',
                    },
                    {
                        'num': 8, 'char': 'ARIA',
                        'visual': 'ARIA to camera; comparison card: CDS used well vs misused; summary bullets; "Next: Swap Pricing" teaser',
                        'onscreen': 'CDS — Summary\n• Transfers credit risk via premium vs contingent payment\n• Spread = market price of default probability\n• 2008 crisis: unhedged naked CDS concentration nearly collapsed the system',
                        'vo': 'A CDS is a powerful credit risk management tool — when used to hedge an actual underlying exposure. The protection buyer transfers default risk for a regular premium. The CDS spread is the market\'s best guess at the probability and severity of default. The 2008 crisis showed what happens when CDS is sold without capital backing and bought without underlying exposure — systemic risk is amplified rather than distributed. Next: how do you price a swap and value it over its life?',
                    },
                ],
            },
        ],
        'next_uor': 'Swap Pricing & Valuation Basics',
    },
    {
        'id': 'SWP.5',
        'title': 'Swap Pricing & Valuation Basics',
        'objective': 'Explain how swaps are priced at inception and how their value changes over time.',
        'competency': 'Swap Valuation Literacy',
        'subcompetencies': [
            'SC5.1 Explain why a swap has zero value at inception',
            'SC5.2 Describe how present value of cash flows determines swap value',
            'SC5.3 Identify factors that cause a swap\'s MTM value to change',
            'SC5.4 Explain the concept of the swap rate',
        ],
        'ears': 'Explain · Describe · Identify',
        'timeline': '10 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Explain (Zero Value at Inception)',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at desk; balance scale graphic — two sides equal at inception',
                        'onscreen': 'Swap Value at Inception = Zero\nFixed leg PV = Floating leg PV\nNeither party is advantaged at the start',
                        'vo': 'Here is something that surprises most people about swaps: at the moment a swap is executed, its fair value to both parties is exactly zero. Why? Because the swap rate — the fixed rate in an IRS — is set so that the present value of the fixed cash flows exactly equals the present value of the expected floating cash flows, using the current yield curve. It is a fair exchange. Nobody is getting a better deal than the other at inception — or they would not agree to it.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; PV calculation diagram — discount factors applied to each cash flow',
                        'onscreen': 'Swap Rate Formula\nSwap Rate = rate where PV(Fixed) = PV(Floating)\nDerived from the yield curve (OIS, LIBOR curve)\nFixing the swap rate makes the trade "fair"',
                        'vo': 'The swap rate is determined by the market — it is the fixed rate that makes the swap fair at inception. To find it, you discount each expected floating payment back to today using discount factors derived from the yield curve. Then you solve for the fixed rate that makes those present values equal. In India, the OIS — Overnight Index Swap — curve is used. The swap rate reflects the market\'s aggregate expectation of future short rates over the swap tenor.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'Animation: time passing, interest rate curve shifting up — one party\'s value increases',
                        'onscreen': 'Value Changes After Inception\nIf rates RISE → Fixed payer gains value (swap MTM positive)\nIf rates FALL → Fixed payer loses value (MTM negative)\nMTM = PV(remaining floating) – PV(remaining fixed)',
                        'vo': 'Once the swap is live, its value to each party moves constantly as interest rates change. If market rates rise above the fixed rate in the swap, the fixed payer is paying less than the going market rate — their position has positive mark-to-market value. Conversely, if rates fall, the fixed payer is locked into a rate above market — negative MTM. The daily valuation — marked to market — drives margining requirements between counterparties under ISDA credit support annexes.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA at screen; list of factors affecting MTM value',
                        'onscreen': 'What Drives Swap MTM?\n1. Changes in interest rates (parallel shifts)\n2. Changes in the yield curve shape (steepening/flattening)\n3. Passage of time (fewer remaining cash flows)\n4. Credit quality of counterparty',
                        'vo': 'Four main factors move a swap\'s mark-to-market value. First, parallel shifts in interest rates — the dominant driver. Second, changes in yield curve shape: if the curve steepens, longer-dated swaps are affected more than short-dated. Third, time decay — as each settlement date passes, there are fewer remaining cash flows to discount. Fourth, counterparty credit quality: if your swap counterparty\'s credit deteriorates, the MTM value of your position is adjusted downward for counterparty risk — called CVA.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Identify & Deeper Explain',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO at desk; numerical MTM example — two cash flows table before and after rate shift',
                        'onscreen': 'MTM Worked Example\nSwap: Pay 7% fixed on Rs 100 Cr, 2 years remaining\nRates rise to 8% → PV(floating remaining) > PV(fixed)\nMTM gain to fixed payer ≈ Rs 1.8 Cr (approx)',
                        'vo': 'Let us make this concrete. You pay 7% fixed, receive floating, Rs 100 crore notional, two years remaining. Rates suddenly rise to 8% across the board. Your fixed leg: two remaining payments of Rs 3.5 crore semi-annually — discounted at the new higher rate. Your floating leg: now expected to pay Rs 4 crore semi-annually — same discount rate. The present value of what you receive exceeds what you owe — your MTM is positive. You could close out the swap and receive a cash payment approximately equal to the PV of that differential.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; unwinding a swap graphic — close-out by entering offsetting swap',
                        'onscreen': 'Exiting a Swap\nOption 1: Enter offsetting swap (pay floating, receive fixed)\nOption 2: Terminate early — pay/receive MTM close-out amount\nOption 3: Assign swap to third party (novation)',
                        'vo': 'Swaps are illiquid compared to exchange-traded derivatives — you cannot simply sell them. To exit, you have three choices. One: enter an exactly offsetting swap — pay floating, receive fixed — which neutralises your net position while leaving both swaps on your books. Two: terminate early with your counterparty and pay or receive the current MTM value — a clean break. Three: novate — transfer your position to a third party with the original counterparty\'s consent. Each approach has different accounting and capital implications.',
                    },
                    {
                        'num': 7, 'char': 'RYO',
                        'visual': 'RYO serious; CVA/DVA concept graphic — credit risk adjustment arrows',
                        'onscreen': 'CVA — Credit Valuation Adjustment\nAdjust MTM for counterparty default risk\nCVA reduces the value of your receivable\nDVA — your own default risk (reverse of CVA)',
                        'vo': 'In post-2008 derivative valuation, a pure rate-based MTM is not sufficient. You must also adjust for the probability that your counterparty defaults before the swap matures — this is the Credit Valuation Adjustment, or CVA. If your counterparty is a shaky bank, the MTM of your positive position is worth less than face value, because you might not collect it. DVA is the mirror image — the adjustment for your own default risk, which creates a perverse situation where your own credit deterioration generates an accounting gain.',
                    },
                    {
                        'num': 8, 'char': 'ARIA',
                        'visual': 'ARIA smiling; summary card with three-point recap; "Next: Swap Applications" teaser',
                        'onscreen': 'Pricing & Valuation — Summary\n• Swap rate makes PV(fixed) = PV(floating) at inception\n• Value = zero at inception; moves with rates thereafter\n• MTM drivers: rate level, curve shape, time, credit quality',
                        'vo': 'To summarise: the swap rate is the fixed rate that makes the trade fair at inception — PV of fixed equals PV of floating. After execution, the swap\'s value fluctuates as interest rates move, time passes, and credit quality changes. Mark-to-market is the mechanism that drives collateral posting under credit support annexes. Understanding valuation is critical for anyone managing a swap book, pricing derivatives for clients, or monitoring counterparty credit exposure. Next: how all of this comes together in real-world applications.',
                    },
                ],
            },
        ],
        'next_uor': 'Swap Applications & Risk Management',
    },
    {
        'id': 'SWP.6',
        'title': 'Swap Applications & Risk Management',
        'objective': 'Identify practical applications of swaps in hedging, liability management, and arbitrage.',
        'competency': 'Swap Applications Literacy',
        'subcompetencies': [
            'SC6.1 Identify how banks use IRS to manage interest rate risk',
            'SC6.2 Explain the comparative advantage argument for currency swaps',
            'SC6.3 Describe how corporates use swaps for liability management',
            'SC6.4 Identify risks in swap positions (counterparty risk, market risk)',
        ],
        'ears': 'Identify · Explain · Describe',
        'timeline': '11 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Identify & Explain',
                'scenes': [
                    {
                        'num': 1, 'char': 'ARIA',
                        'visual': 'ARIA at desk; bank balance sheet graphic — asset/liability mismatch highlighted',
                        'onscreen': 'Banks & IRS — Managing the Balance Sheet\nAssets: long-term fixed rate loans\nLiabilities: short-term floating deposits\nMismatch creates interest rate risk → hedge with IRS',
                        'vo': 'Banks are natural users of interest rate swaps because their balance sheets are structurally mismatched. Loans to corporates are often five to ten year fixed rate. But deposits funding those loans reprice every few months. If rates rise, the bank pays more on deposits but earns the same on loans — a squeeze on net interest margin. Banks enter receive-fixed IRS to hedge this mismatch. They receive fixed from the swap, offsetting their fixed-rate loans, while the floating payment on the swap matches their floating deposit costs.',
                    },
                    {
                        'num': 2, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; two-company comparative advantage table with borrowing costs and savings',
                        'onscreen': 'Comparative Advantage — Currency Swap\nCompany A: better in INR markets (8% vs 10%)\nCompany B: better in USD markets (5% vs 7%)\nTotal saving from swapping: 2% to share between them',
                        'vo': 'The comparative advantage argument for currency swaps is elegant and powerful. Company A can borrow INR at 8% or USD at 6%. Company B can borrow INR at 10% or USD at 5%. Company A has comparative advantage in INR; Company B in USD. If each borrows where they are relatively better, and then swaps, the total cost saving is 2% per annum on the notional — to be divided between the two companies. This is arbitrage of different credit standings across different markets.',
                    },
                    {
                        'num': 3, 'char': 'ARIA',
                        'visual': 'ARIA at screen; corporate liability management timeline — debt restructuring via swap',
                        'onscreen': 'Corporate Liability Management\nConvert fixed debt → floating (if rates expected to fall)\nConvert floating → fixed (if rates expected to rise)\nExtend or shorten duration without refinancing',
                        'vo': 'Swaps allow corporates to manage the structure of their liabilities without the cost and disruption of refinancing. A company with Rs 500 crore of 8% fixed-rate bonds maturing in 5 years can enter an IRS — pay floating, receive 8% fixed — effectively converting their liability to floating. If market rates subsequently fall to 6%, they now effectively pay only 6% on that debt. This is far cheaper and faster than calling the bonds and reissuing at the lower rate.',
                    },
                    {
                        'num': 4, 'char': 'RYO',
                        'visual': 'RYO at desk; risk spectrum graphic — counterparty risk, market risk, liquidity risk, operational risk',
                        'onscreen': 'Risks in Swap Positions\n1. Counterparty (Credit) Risk — largest for long-dated swaps\n2. Market Risk — MTM moves with rates/FX\n3. Liquidity Risk — OTC, harder to exit\n4. Operational Risk — settlement, documentation',
                        'vo': 'Swap positions carry multiple risk dimensions. Counterparty risk is the most significant for long-dated OTC swaps — the risk that your counterparty defaults before the swap matures. Market risk is the day-to-day MTM movement from interest rate and exchange rate changes. Liquidity risk arises from the OTC nature — you cannot simply sell a swap like a bond; unwinding requires an offsetting trade or termination. Operational risk covers settlement failures, ISDA documentation errors, and system failures.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Describe & Identify (Risk Management)',
                'scenes': [
                    {
                        'num': 5, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; hedge accounting graphic — IFRS 9 fair value hedge vs cash flow hedge',
                        'onscreen': 'Hedge Accounting for Swaps\nFair Value Hedge: hedge of fixed-rate asset/liability\nCash Flow Hedge: hedge of variable cash flows\nIFRS 9 / Ind AS 109 requirements',
                        'vo': 'When corporates use swaps as hedges, they must decide on hedge accounting treatment under IFRS 9 or Indian Accounting Standard 109. A fair value hedge is used when hedging a fixed-rate asset or liability — changes in the hedged item\'s fair value and the swap\'s fair value offset each other in P&L. A cash flow hedge is used for floating rate exposures — the effective portion of the swap\'s gain or loss goes to Other Comprehensive Income, not P&L, until the hedged cash flow occurs. Incorrect classification creates P&L volatility.',
                    },
                    {
                        'num': 6, 'char': 'RYO',
                        'visual': 'RYO at screen; central clearing graphic — CCP standing between two swap counterparties',
                        'onscreen': 'Central Clearing for Swaps\nCCP becomes counterparty to both sides\nEliminate bilateral counterparty risk\nMargin posted daily — variation margin + initial margin\nMandatory for standardised IRS under EMIR / Dodd-Frank',
                        'vo': 'Post-2008 regulations require standardised IRS contracts to be cleared through a central counterparty — a CCP. The CCP steps in as buyer to every seller and seller to every buyer. This eliminates bilateral counterparty risk — you no longer care if your original counterparty defaults, because the CCP is your counterparty now. In return, you must post initial margin and daily variation margin with the CCP. In India, CCIL — the Clearing Corporation of India — provides central clearing for OTC derivatives including IRS and OIS.',
                    },
                    {
                        'num': 7, 'char': 'ARIA',
                        'visual': 'ARIA serious; speculative use graphic — macro hedge fund taking large IRS positions on RBI rate view',
                        'onscreen': 'Speculation with Swaps\nMacro funds: express rate views via IRS\nPay fixed = bet rates will RISE\nReceive fixed = bet rates will FALL\nLevered exposure without bond purchase',
                        'vo': 'Not all swap users are hedgers. Macro hedge funds routinely use IRS as a leveraged way to express interest rate views. If a fund believes the RBI will raise rates aggressively, they pay fixed in an OIS — they receive floating, which will grow as rates rise. This gives them rate exposure equivalent to a large short bond position, but without the financing cost of actually borrowing and shorting bonds. The leverage is embedded in the notional — a Rs 1,000 crore notional swap requires only margin, not the full Rs 1,000 crore.',
                    },
                    {
                        'num': 8, 'char': 'RYO',
                        'visual': 'RYO to camera; final summary card; module completion animation',
                        'onscreen': 'Module SWP Complete\nSwaps: OTC agreements to exchange cash flows\nIRS, Currency Swaps, CDS — each addresses a different risk\nApplications: hedging, liability management, arbitrage, speculation',
                        'vo': 'You have now covered the full sweep of the Derivatives: SWAPS module. From the foundational concept of a swap as a cash-flow exchange, through the mechanics of interest rate and currency swaps, the credit risk dimension of CDS, the valuation framework, and finally the real-world applications in banking, corporate treasury, and speculation. Swaps are among the most powerful and most widely used instruments in global capital markets — and understanding them is essential for any role in investment banking operations, treasury, or risk management.',
                    },
                ],
            },
        ],
        'next_uor': None,  # Module complete
    },
]

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    elif level == 2:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)
    return p

def add_mapping_table(doc, uor):
    # Header
    p = doc.add_paragraph()
    run = p.add_run('Internal Mapping Table')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    # Column widths (approx)
    widths = [Inches(0.9), Inches(1.4), Inches(2.2), Inches(1.0), Inches(1.8)]
    headers = ['UOR ID', 'UOR Title', 'Learning Objectives', 'EARs', 'Sub-competencies']
    header_row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = header_row.cells[i]
        cell.width = w
        set_cell_bg(cell, '1E3A5F')
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(hdr)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(9)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data_row = table.rows[1]
    vals = [
        uor['id'],
        uor['title'],
        uor['objective'],
        uor['ears'],
        '\n'.join(uor['subcompetencies'])
    ]
    for i, (val, w) in enumerate(zip(vals, widths)):
        cell = data_row.cells[i]
        cell.width = w
        set_cell_bg(cell, 'EEF4FB')
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(val)
        run2.font.size = Pt(9)

    doc.add_paragraph()  # spacer

def add_script_table(doc, video):
    p = doc.add_paragraph()
    run = p.add_run(f'Production Script — {video["label"]}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    cols = ['Scene #', 'Character', 'Visual Cue', 'On-Screen Text', 'Voice Over']
    col_widths = [Inches(0.5), Inches(0.7), Inches(1.5), Inches(1.6), Inches(3.0)]
    n_scenes = len(video['scenes'])
    table = doc.add_table(rows=n_scenes + 1, cols=5)
    table.style = 'Table Grid'

    # Header row
    hrow = table.rows[0]
    for i, (col, w) in enumerate(zip(cols, col_widths)):
        cell = hrow.cells[i]
        cell.width = w
        set_cell_bg(cell, '005F5F')
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(col)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(9)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    alt_colors = ['FFFFFF', 'F0F7F4']
    for r_idx, scene in enumerate(video['scenes']):
        row = table.rows[r_idx + 1]
        fill = alt_colors[r_idx % 2]
        vals = [
            str(scene['num']),
            scene['char'],
            scene['visual'],
            scene['onscreen'],
            scene['vo']
        ]
        for i, (val, w) in enumerate(zip(vals, col_widths)):
            cell = row.cells[i]
            cell.width = w
            set_cell_bg(cell, fill)
            p2 = cell.paragraphs[0]
            run2 = p2.add_run(val)
            run2.font.size = Pt(8.5)
            if i == 1:  # Character name bold
                run2.bold = True
                if val == 'RYO':
                    run2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
                else:
                    run2.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)

    doc.add_paragraph()  # spacer

def add_footer(doc, text):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─────────────────────────────────────────────
# BUILD VIDEO PRODUCTION DOCUMENT
# ─────────────────────────────────────────────
def build_video_doc():
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DERIVATIVES: SWAPS')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Video Production Document')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('CIBOP Module 7 | Version 2')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('Characters: RYO (Male, Navy Suit, Mentor, Dry Wit) | ARIA (Female, Teal Blazer, Analyst, Sharp Comic Timing)')
    run4.font.size = Pt(10)
    run4.italic = True

    doc.add_paragraph()
    doc.add_page_break()

    # TOC note
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run('TABLE OF CONTENTS')
    run_toc.bold = True
    run_toc.font.size = Pt(12)
    run_toc.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    for uor in UORS:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.3)
        run_item = p_item.add_run(f'{uor["id"]}  {uor["title"]}')
        run_item.font.size = Pt(10)

    doc.add_page_break()

    # Per-UOR content
    for uor_idx, uor in enumerate(UORS):
        # UOR Header
        add_heading(doc, f'{uor["id"]} — {uor["title"]}', level=1)
        p_tag = doc.add_paragraph()
        run_tag = p_tag.add_run(f'EARs: {uor["ears"]}  |  Competency: {uor["competency"]}  |  Timeline: {uor["timeline"]}')
        run_tag.font.size = Pt(9)
        run_tag.italic = True
        run_tag.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()

        # Internal Mapping Table
        add_mapping_table(doc, uor)

        # Video scripts
        for video in uor['videos']:
            add_script_table(doc, video)

        # Transition card
        next_title = uor['next_uor']
        if next_title:
            trans_text = f'End of {uor["title"]} — Next: {next_title}'
        else:
            trans_text = f'End of {uor["title"]} — Module Complete'

        p_trans = doc.add_paragraph()
        p_trans.paragraph_format.space_before = Pt(12)
        p_trans.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Styled box via shading on a table
        trans_table = doc.add_table(rows=1, cols=1)
        trans_table.style = 'Table Grid'
        tc = trans_table.rows[0].cells[0]
        set_cell_bg(tc, '1E3A5F')
        p_tc = tc.paragraphs[0]
        p_tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tc = p_tc.add_run(f'▶  {trans_text}')
        run_tc.bold = True
        run_tc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_tc.font.size = Pt(11)

        if uor_idx < len(UORS) - 1:
            doc.add_page_break()

    # Footer
    add_footer(doc, 'CIBOP PPT 07 — Derivatives: SWAPS  |  Confidential — Internal Use Only')

    path = OUTPUT_DIR + 'SWP_Video_Production_Document_v2.docx'
    doc.save(path)
    print(f'Saved: {path}')
    return path

if __name__ == '__main__':
    p = build_video_doc()
    import os
    print(f'File size: {os.path.getsize(p):,} bytes')
    print('Video Production Document DONE.')
