"""Export — download approved content as formatted DOCX files."""

import re
import io
import streamlit as st
from lib.db import get_generated, get_plan_items, get_review, get_all_audits_for_topic
from lib.sidebar import render_topic_selector

st.set_page_config(page_title="Export — CIBOP", layout="wide")
render_topic_selector()
st.title("📥 Export")

if "selected_topic_id" not in st.session_state:
    st.warning("Please select a topic first.")
    st.stop()

topic_id = st.session_state["selected_topic_id"]
topic_name = st.session_state.get("selected_topic_name", "")
module_code = st.session_state.get("selected_topic_code", "CIBOP")
st.subheader(f"Module: {topic_name}")

generated = get_generated(topic_id)
plan_items = get_plan_items(topic_id)
audit_data = {a["content"]["id"]: a["audit"] for a in get_all_audits_for_topic(topic_id) if a["audit"]}

if not generated:
    st.warning("No content generated yet.")
    st.stop()

# ── Export options ─────────────────────────────────────────────────────────────
st.markdown("### Export Options")
include_unreviewed = st.checkbox("Include content not yet reviewed", value=False)
include_failed     = st.checkbox("Include content that failed audit (<80%)", value=False)
st.markdown("---")


# ══════════════════════════════════════════════════════════════════════════════
# DOCX helpers
# ══════════════════════════════════════════════════════════════════════════════

def _parse_script_table(script_text: str):
    """Parse the markdown table in a video script.
    Returns (header_cells, data_rows) where each row is a list of strings."""
    lines = script_text.split("\n")
    header = None
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [c.strip() for c in stripped.split("|")[1:-1]]
        if not cells:
            continue
        if all(re.fullmatch(r"[-: ]+", c) for c in cells):
            continue  # separator row
        if header is None:
            header = cells
        else:
            while len(cells) < len(header):
                cells.append("")
            rows.append(cells[:len(header)])
    return header or [], rows


def _slide_refs(script_text: str) -> str:
    m = re.search(r"SLIDE_REFS_USED:\s*(.+)", script_text)
    return m.group(1).strip() if m else ""


def _set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_divider(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_before = Pt(6)
    div_p.paragraph_format.space_after  = Pt(6)
    pPr = div_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_script_table(doc, script_text: str, audit: dict | None):
    """Add a formatted production script table to the document."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
    TEAL  = RGBColor(0x00, 0x5F, 0x5F)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    # Character background / foreground colours (hex bg, RGBColor fg)
    CHAR_STYLE = {
        "MOTION": ("E8F0FA", NAVY),
        "RYO":    ("E8FAF5", TEAL),
        "ARIA":   ("F5EAF5", RGBColor(0x6B, 0x3A, 0x8B)),
        "BOTH":   ("FAF0E8", RGBColor(0x8B, 0x45, 0x13)),
    }

    # Column widths in cm, total = 18 cm (A4 minus 1.5 cm margins each side)
    COL_CM = [0.5, 1.4, 4.2, 3.5, 8.4]

    header_cells, rows = _parse_script_table(script_text)
    if not header_cells or not rows:
        doc.add_paragraph(script_text, style="Normal")
        return

    # Audit score line above table
    if audit:
        cov  = audit.get("coverage_score", 0)
        ord_ = audit.get("order_score", 0)
        fid  = audit.get("fidelity_score", 0)
        overall = (cov + ord_ + fid) / 3
        sp = doc.add_paragraph()
        sr = sp.add_run(
            f"Audit: {overall:.0f}%  |  Coverage {cov:.0f}%  |  "
            f"Order {ord_:.0f}%  |  Fidelity {fid:.0f}%"
        )
        sr.font.size = Pt(9)
        sr.font.italic = True
        sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        sp.paragraph_format.space_after = Pt(4)

    # Build the table (header row + data rows)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    tbl.style = "Table Grid"

    # Enforce column widths
    for col_idx, width_cm in enumerate(COL_CM[:len(header_cells)]):
        for tbl_row in tbl.rows:
            tbl_row.cells[col_idx].width = Cm(width_cm)

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(header_cells):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, "1E3A5F")
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Detect which column is "Character"
    char_col = next(
        (i for i, h in enumerate(header_cells) if "character" in h.lower()), 1
    )

    # Data rows
    for r_idx, row_vals in enumerate(rows):
        tbl_row = tbl.rows[r_idx + 1]
        char_val = row_vals[char_col].upper() if len(row_vals) > char_col else ""
        char_bg, char_fg = CHAR_STYLE.get(char_val, ("FFFFFF", RGBColor(0x22, 0x22, 0x22)))
        row_bg = "F8F8F8" if r_idx % 2 else "FFFFFF"

        for c_idx, cell_text in enumerate(row_vals[:len(header_cells)]):
            cell = tbl_row.cells[c_idx]
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            is_char = (c_idx == char_col)
            _set_cell_bg(cell, char_bg if is_char else row_bg)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if is_char:
                run.font.bold = True
                run.font.color.rgb = char_fg
            else:
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Slide refs
    refs = _slide_refs(script_text)
    if refs:
        rp = doc.add_paragraph()
        rr = rp.add_run(f"📎  Slides referenced: {refs}")
        rr.font.size = Pt(8)
        rr.font.italic = True
        rr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        rp.paragraph_format.space_before = Pt(3)
        rp.paragraph_format.space_after  = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# Video DOCX builder
# ══════════════════════════════════════════════════════════════════════════════

def build_video_docx(items, plan_map, audit_map, include_unreviewed, include_failed):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1E, 0x3A, 0x5F)
    TEAL = RGBColor(0x00, 0x5F, 0x5F)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(1.5))

    # Cover title
    tp = doc.add_paragraph()
    tr = tp.add_run(f"{topic_name} — Video Production Document")
    tr.font.size = Pt(18)
    tr.font.color.rgb = NAVY
    tr.font.bold = True
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(20)

    current_uor = None

    for item in sorted(items, key=lambda x: (x["uor_id"], x["sc_id"])):
        if item["content_type"] != "video_script":
            continue

        review = get_review(item["id"])
        audit  = audit_map.get(item["id"])

        if not include_unreviewed and not review:
            continue
        if not include_failed and audit:
            ov = (audit.get("coverage_score", 0) + audit.get("order_score", 0) +
                  audit.get("fidelity_score", 0)) / 3
            if ov < 80:
                continue

        # UOR header (new page per UOR)
        if item["uor_id"] != current_uor:
            if current_uor is not None:
                doc.add_page_break()
            current_uor = item["uor_id"]
            plan_item = plan_map.get(f"{item['uor_id']}_{item['sc_id']}")
            uor_title = plan_item.get("uor_title", "") if plan_item else ""

            up = doc.add_paragraph()
            ur = up.add_run(f"UOR {item['uor_id']}: {uor_title}")
            ur.font.size = Pt(14)
            ur.font.bold = True
            ur.font.color.rgb = NAVY
            up.paragraph_format.space_after = Pt(8)

        # SC header
        sp = doc.add_paragraph()
        sr = sp.add_run(f"🎬  {item['sc_id']} — Video Script")
        sr.font.size = Pt(12)
        sr.font.bold = True
        sr.font.color.rgb = TEAL
        sp.paragraph_format.space_before = Pt(12)
        sp.paragraph_format.space_after  = Pt(4)

        # Script table
        content = review["edited_content"] if review else item["content_text"]
        _add_script_table(doc, content, audit)

        _add_divider(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# Assessment DOCX builder
# ══════════════════════════════════════════════════════════════════════════════

def build_assessment_docx(items, plan_map, audit_map, include_unreviewed, include_failed):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    NAVY = RGBColor(0x1E, 0x3A, 0x5F)
    TEAL = RGBColor(0x00, 0x5F, 0x5F)
    GREY = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(1.5))

    tp = doc.add_paragraph()
    tr = tp.add_run(f"{topic_name} — Knowledge Assessment")
    tr.font.size = Pt(18)
    tr.font.color.rgb = NAVY
    tr.font.bold = True
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(20)

    q_num = 1
    current_uor = None

    for item in sorted(items, key=lambda x: (x["uor_id"], x["sc_id"])):
        if item["content_type"] != "assessment":
            continue

        review = get_review(item["id"])
        audit  = audit_map.get(item["id"])

        if not include_unreviewed and not review:
            continue
        if not include_failed and audit:
            ov = (audit.get("coverage_score", 0) + audit.get("order_score", 0) +
                  audit.get("fidelity_score", 0)) / 3
            if ov < 80:
                continue

        if item["uor_id"] != current_uor:
            current_uor = item["uor_id"]
            plan_item = plan_map.get(f"{item['uor_id']}_{item['sc_id']}")
            uor_title = plan_item.get("uor_title", "") if plan_item else ""
            doc.add_paragraph()
            up = doc.add_paragraph()
            ur = up.add_run(f"UOR {item['uor_id']}: {uor_title}")
            ur.font.size = Pt(12)
            ur.font.bold = True
            ur.font.color.rgb = NAVY
            up.paragraph_format.space_after = Pt(6)

        qh = doc.add_paragraph()
        qr = qh.add_run(f"Q{q_num}  |  {item['uor_id']} / {item['sc_id']}")
        qr.font.size = Pt(11)
        qr.font.bold = True
        qr.font.color.rgb = TEAL
        qh.paragraph_format.space_before = Pt(10)
        qh.paragraph_format.space_after  = Pt(4)
        q_num += 1

        if audit:
            cov = audit.get("coverage_score", 0)
            ord_ = audit.get("order_score", 0)
            fid  = audit.get("fidelity_score", 0)
            ov = (cov + ord_ + fid) / 3
            asp = doc.add_paragraph()
            asr = asp.add_run(
                f"Audit {ov:.0f}%  |  Coverage {cov:.0f}%  |  Order {ord_:.0f}%  |  Fidelity {fid:.0f}%"
            )
            asr.font.size = Pt(8)
            asr.font.italic = True
            asr.font.color.rgb = GREY
            asp.paragraph_format.space_after = Pt(3)

        content = review["edited_content"] if review else item["content_text"]
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped or stripped.startswith("SLIDE_REFS_USED"):
                continue
            p = doc.add_paragraph(stripped)
            if p.runs:
                p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)

        _add_divider(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# Export buttons
# ══════════════════════════════════════════════════════════════════════════════

plan_map = {f"{i['uor_id']}_{i['sc_id']}": i for i in plan_items}

col1, col2 = st.columns(2)

with col1:
    st.markdown("#### 📹 Video Production Document")
    video_items = [g for g in generated if g["content_type"] == "video_script"]
    st.caption(f"{len(video_items)} video scripts")
    if st.button("Build Video DOCX", type="primary"):
        with st.spinner("Building document…"):
            try:
                buf = build_video_docx(generated, plan_map, audit_data,
                                       include_unreviewed, include_failed)
                st.download_button(
                    label="⬇️ Download Video Production DOCX",
                    data=buf,
                    file_name=f"{topic_name.replace(' ', '_')}_Video_Production.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                import traceback
                st.error(f"Build failed: {e}")
                st.code(traceback.format_exc())

with col2:
    st.markdown("#### 📝 Knowledge Assessment")
    assess_items = [g for g in generated if g["content_type"] == "assessment"]
    st.caption(f"{len(assess_items)} assessment questions")
    if st.button("Build Assessment DOCX", type="primary"):
        with st.spinner("Building document…"):
            try:
                buf = build_assessment_docx(generated, plan_map, audit_data,
                                            include_unreviewed, include_failed)
                st.download_button(
                    label="⬇️ Download Assessment DOCX",
                    data=buf,
                    file_name=f"{topic_name.replace(' ', '_')}_Assessment.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                import traceback
                st.error(f"Build failed: {e}")
                st.code(traceback.format_exc())


# ── Status summary ─────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown("### Content Status Summary")

data = []
for g in generated:
    review  = get_review(g["id"])
    audit   = audit_data.get(g["id"])
    overall = None
    if audit:
        overall = (audit.get("coverage_score", 0) + audit.get("order_score", 0) +
                   audit.get("fidelity_score", 0)) / 3
    data.append({
        "UOR":         g["uor_id"],
        "SC":          g["sc_id"],
        "Type":        "📹 Video" if g["content_type"] == "video_script" else "📝 Assessment",
        "Audit Score": f"{overall:.0f}%" if overall is not None else "—",
        "Pass":        "✅" if overall and overall >= 80 else ("❌" if overall else "—"),
        "Review":      ("✅ Approved" if (review and review.get("approved"))
                        else ("❌ Rejected" if (review and not review.get("approved"))
                              else "⏳ Pending")),
    })

if data:
    import pandas as pd
    st.dataframe(pd.DataFrame(data), use_container_width=True, hide_index=True)
