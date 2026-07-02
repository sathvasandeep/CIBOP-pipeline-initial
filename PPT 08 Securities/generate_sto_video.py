#!/usr/bin/env python3
"""Generate STO & RDM Video Production Document — CIBOP Module 8"""
import subprocess, sys
subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'], check=True)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = '/Users/sandeeppr/CIBOP/PPT 08 Securities/STO_Video_Production_Document_v2.docx'

# ─────────────────────────────────────────────
# UOR DATA — Securities Trading Organisation & RDM
# ─────────────────────────────────────────────
UORS = [
    {
        'id': 'STO.1',
        'title': 'Introduction to the Securities Trading Organisation',
        'objective': 'Describe the structure of a Securities Trading Organisation and the roles of the front, middle, and back office.',
        'competency': 'STO Structure & Function',
        'subcompetencies': [
            'SC1.1 Define the Securities Trading Organisation (STO) and its purpose',
            'SC1.2 Identify the three primary divisions: Front Office, Middle Office, Back Office',
            'SC1.3 Explain the flow of a trade from execution to settlement through the STO',
        ],
        'ears': 'Define · Identify · Explain',
        'timeline': '10 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Define & Identify',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at desk; wide graphic of an investment bank floor — three labelled zones: Front Office (green), Middle Office (amber), Back Office (blue)',
                        'onscreen': 'The Securities Trading Organisation\nThe engine behind every trade in capital markets\nFront • Middle • Back — each with a distinct role',
                        'vo': 'Welcome to Module 8. Before a single trade can settle — before money moves, before securities change hands — an entire organisation has to function flawlessly. The Securities Trading Organisation, or STO, is that machine. It is divided into three main divisions: the front office, which executes and manages trades; the middle office, which monitors risk and ensures compliance; and the back office, which confirms, clears, and settles. Think of it as the three floors of a bank — glamorous upstairs, essential downstairs, and the risk police in between.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; org chart with three columns — Front Office roles, Middle Office roles, Back Office roles',
                        'onscreen': 'Front Office\nTraders · Sales · Research · Structuring\n\nMiddle Office\nRisk · P&L Attribution · Compliance\n\nBack Office\nOperations · Settlement · Reconciliation · Custody',
                        'vo': 'Let us populate each division. The front office includes traders, sales staff, research analysts, and structurers — these are the revenue-generating roles. The middle office provides the governance layer: risk managers who monitor exposure in real time, P&L attribution teams who verify the front office\'s reported profits, and compliance officers who ensure trades meet regulatory requirements. The back office handles the operational mechanics: confirming trades, instructing settlement, managing custody, and reconciling positions. No settlement without the back office.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'RYO pointing at trade lifecycle flow diagram — arrow from Trade Execution → Trade Capture → Confirmation → Clearing → Settlement',
                        'onscreen': 'Trade Lifecycle in the STO\nExecution → Capture → Confirmation → Clearing → Settlement\nEach stage owned by a different part of the STO',
                        'vo': 'Every trade follows the same lifecycle through the STO. Execution happens in the front office — the trader hits the button. The trade is then captured in the trade management system. Confirmation is sent to the counterparty — both sides must agree on the terms. The trade goes to clearing — either through a central counterparty or bilaterally. Finally, settlement occurs — securities and cash actually change hands. Fail at any stage, and you have a settlement failure, a fine, or worse. The back office owns steps three through five.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA to camera; Chinese Wall graphic separating front office from middle/back office',
                        'onscreen': 'The Chinese Wall\nStrict information barrier between Front and Back Office\nPrevents insider trading and front-running\nRegulatory requirement — not just good practice',
                        'vo': 'One of the most important structural features of an STO is the Chinese Wall — a strict information barrier separating the front office from the rest. The back office handles client information, settlement details, and upcoming corporate actions. If that information leaked to traders, they could front-run orders or insider trade. Regulators require this separation, and firms enforce it through physical separation, system access controls, and email monitoring. Crossing the Chinese Wall, even accidentally, is a regulatory event.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Explain (Trade Flow & Systems)',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO at screen; systems map — OMS, EMS, Trade Capture system, Risk system, Settlement system connected by arrows',
                        'onscreen': 'STO Systems Landscape\nOMS: Order Management System\nEMS: Execution Management System\nTrade Capture → Risk → Settlement systems\nData flows through each in sequence',
                        'vo': 'The STO runs on interconnected technology systems. Orders originate in the OMS — Order Management System — where they are created, routed, and tracked. Execution happens via the EMS — Execution Management System — which connects to exchanges and dark pools. Post-execution, the trade flows into the trade capture system, where it is booked. The risk system reads this in real time. Then operations teams use the settlement system to instruct custodians and clearinghouses. Each system must talk to the next — and a break in that chain means manual intervention, which means risk.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA at desk; straight-through processing (STP) graphic — trade flowing automatically from execution to settlement with no human touches',
                        'onscreen': 'Straight-Through Processing (STP)\nIdeal: zero human touches from execution to settlement\nReality: exceptions require manual intervention\nSTP rate = key operational efficiency metric',
                        'vo': 'The holy grail of back office operations is Straight-Through Processing, or STP. In a fully STP environment, a trade executed by the front office flows automatically — without human intervention — all the way through to settlement. The system confirms, clears, and settles without anyone touching a keyboard. In practice, exceptions break this flow: unmatched trades, static data errors, insufficient cash. The STP rate — the percentage of trades that settle without manual intervention — is a critical KPI for any operations team. High STP rate: lower cost, lower risk.',
                    },
                    {
                        'num': 7, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; key STO metrics — settlement fail rate, STP rate, reconciliation breaks, cost per trade',
                        'onscreen': 'STO Performance Metrics\n• Settlement fail rate (target: <1%)\n• STP rate (target: >95%)\n• Reconciliation breaks (target: zero aged items)\n• Cost per trade (efficiency benchmark)',
                        'vo': 'How do you measure whether an STO is performing well? Four metrics matter most. Settlement fail rate: what percentage of trades fail to settle on value date — anything above 1% draws regulatory scrutiny. STP rate: the inverse of manual intervention — industry best practice is above 95%. Reconciliation breaks: unresolved differences between internal records and external statements — aged breaks are a red flag for auditors. And cost per trade: the total operational cost divided by trade volume — a metric that drives investment in automation.',
                    },
                    {
                        'num': 8, 'char': 'ARIA',
                        'visual': 'ARIA smiling; summary card; "Next: Front Office in Detail" teaser',
                        'onscreen': 'STO Summary\n• Three divisions: Front, Middle, Back Office\n• Trade lifecycle: Execution → Settlement\n• Chinese Wall: regulatory requirement\n• STP rate: the efficiency benchmark',
                        'vo': 'The Securities Trading Organisation is a precisely engineered machine in which every role, every system, and every process has a purpose. The front office generates revenue. The middle office controls risk. The back office ensures that trades actually settle — that securities and cash move correctly and on time. In the sessions that follow, we will drill into each division and the critical function of reference data that underpins everything. Next: the front office in detail.',
                    },
                ],
            },
        ],
        'next_uor': 'Front Office Operations',
    },
    {
        'id': 'STO.2',
        'title': 'Front Office Operations',
        'objective': 'Identify the key roles and responsibilities of the front office in trade execution and client interaction.',
        'competency': 'Front Office Function',
        'subcompetencies': [
            'SC2.1 Identify front office roles: trader, sales, research, structuring',
            'SC2.2 Explain the order flow from client instruction to trade execution',
            'SC2.3 Describe the use of OMS and EMS in front office execution',
            'SC2.4 Explain pre-trade compliance checks and best execution obligations',
        ],
        'ears': 'Identify · Explain · Describe',
        'timeline': '11 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Identify & Explain',
                'scenes': [
                    {
                        'num': 1, 'char': 'ARIA',
                        'visual': 'ARIA at desk; busy trading floor graphic — screens with live prices, traders on phones',
                        'onscreen': 'The Front Office\nRevenue-generating engine of the STO\nTraders · Sales · Research · Structuring',
                        'vo': 'The front office is where the action is. It is the revenue-generating engine of the investment bank — the part you see in every Wall Street film, the part that gets the bonuses. But it is also the part that bears direct market risk. The front office is made up of four main groups: traders who take positions and execute; sales staff who manage client relationships and channel client orders to traders; research analysts who generate investment ideas and market intelligence; and structurers who design bespoke product solutions for clients.',
                    },
                    {
                        'num': 2, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; order flow diagram: Client → Sales → Trader → Exchange/Market → Trade Capture',
                        'onscreen': 'Order Flow in the Front Office\nClient calls Sales with order\nSales passes to Trader\nTrader executes on market\nTrade captured in OMS',
                        'vo': 'Here is a typical order flow. A fund manager wants to buy ten million shares of Reliance Industries. She calls her sales coverage at the bank. Sales logs the order and passes it — with any specific instructions like limit price or execution algo — to the equity trader. The trader routes the order through the OMS to the exchange via the EMS. Once filled, the execution details are captured in the OMS and flow downstream to operations. The whole process can take seconds. Getting any step wrong — wrong ISIN, wrong side, wrong quantity — is called a trade break, and someone\'s afternoon just got very long.',
                    },
                    {
                        'num': 3, 'char': 'ARIA',
                        'visual': 'ARIA at screen; OMS screenshot showing order book, algo selection, pre-trade compliance check overlay',
                        'onscreen': 'OMS — Order Management System\n• Creates and tracks orders from receipt to fill\n• Pre-trade compliance checks (position limits, restricted list)\n• Algo selection: VWAP, TWAP, Implementation Shortfall\n• Feeds to EMS for market execution',
                        'vo': 'The Order Management System is the front office\'s command centre. Every order is created here, checked here, and tracked here. Before an order is released to market, the OMS runs automated pre-trade compliance checks: is this security on the restricted list? Does this trade breach position limits? Is the client\'s account approved for this instrument? These checks happen in milliseconds. If an order passes, it is routed to the EMS where execution algorithms — VWAP, TWAP, Implementation Shortfall — determine how to work the order in the market.',
                    },
                    {
                        'num': 4, 'char': 'RYO',
                        'visual': 'RYO to camera; best execution policy graphic — MiFID II / SEBI requirements',
                        'onscreen': 'Best Execution Obligation\nFirms must achieve the best possible outcome for clients\nFactors: price, speed, likelihood of execution, size, cost\nDocumented in Best Execution Policy — regulatory requirement',
                        'vo': 'Best execution is a regulatory obligation, not a courtesy. Under MiFID II in Europe and equivalent SEBI regulations in India, firms must take all sufficient steps to achieve the best possible result for clients when executing orders. Best is not just about price — it also considers execution speed, likelihood of execution, the size of the order, and implicit costs like market impact. Firms must maintain a Best Execution Policy, document how orders were executed, and be able to demonstrate compliance on request. Getting this wrong is an enforcement action.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Describe (Roles, Risk, P&L)',
                'scenes': [
                    {
                        'num': 5, 'char': 'ARIA',
                        'visual': 'ARIA at desk; proprietary trading vs agency trading distinction graphic',
                        'onscreen': 'Proprietary vs Agency Trading\nAgency: bank executes on behalf of client (no market risk)\nProprietary: bank trades own capital (bears market risk)\nProp trading restricted post-2008 (Volcker Rule)',
                        'vo': 'There are two modes of trading in the front office. Agency trading: the bank executes a client\'s order, taking no market risk — it earns a commission. Proprietary trading: the bank uses its own capital to take positions, bearing full market risk in exchange for profit potential. Post the 2008 financial crisis, the Volcker Rule in the US and similar measures globally sharply restricted bank proprietary trading. Most banks now separate or have divested their prop trading desks. The distinction matters because it fundamentally changes who bears the risk.',
                    },
                    {
                        'num': 6, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; P&L attribution breakdown — realised P&L, unrealised MTM, fees & commissions, funding costs',
                        'onscreen': 'Front Office P&L\nRealised P&L: profit from closed positions\nUnrealised (MTM): open positions marked to market\nNet of: commissions, spreads, financing costs\nMiddle office verifies independently',
                        'vo': 'Every front office desk produces a daily P&L — profit and loss. This is made up of realised profit from positions that have been closed, plus unrealised mark-to-market gains or losses on open positions. Net of commissions paid to brokers, bid-offer spreads, and financing costs on borrowed securities or leveraged positions. The critical governance point: the middle office independently verifies the front office P&L. Traders do not mark their own books — the middle office uses independent pricing sources to check. Any discrepancy triggers an escalation.',
                    },
                    {
                        'num': 7, 'char': 'ARIA',
                        'visual': 'ARIA at screen; errors and omissions in trade capture — wrong ISIN, wrong side, duplicate trade examples',
                        'onscreen': 'Front Office Trade Errors\nWrong ISIN / wrong security\nWrong side (buy vs sell — most expensive error)\nWrong quantity or price\nDuplicate trade\nAll require immediate correction and escalation',
                        'vo': 'Front office errors are expensive. The most common: wrong security ISIN — a buy of BHP when you meant BP. Wrong side — buying instead of selling — this can move markets if the size is large and costs the firm the full round-trip spread plus market impact. Wrong quantity — particularly dangerous in illiquid markets. Duplicate trades — the same order entered twice, doubling the exposure. All errors must be captured, corrected, and reported. Most firms have an error policy requiring escalation to risk management for anything above a threshold — typically £10,000 to £100,000 P&L impact.',
                    },
                    {
                        'num': 8, 'char': 'RYO',
                        'visual': 'RYO smiling; summary card with four key takeaways; "Next: Middle Office Operations" teaser',
                        'onscreen': 'Front Office Summary\n• Four roles: traders, sales, research, structuring\n• Order flow: client → sales → trader → market → capture\n• Best execution: legal obligation, documented policy\n• P&L independently verified by middle office',
                        'vo': 'The front office is the sharp end of the bank — revenue, risk, and real-time decision making. Its roles are distinct: traders execute, sales manage relationships, research generates ideas, structurers create solutions. Orders flow from client instruction through systematic pre-trade checks to market execution and trade capture. Best execution is a regulatory obligation, and P&L is independently verified. Next: we move one floor down to the middle office — the risk and governance layer that keeps the front office honest.',
                    },
                ],
            },
        ],
        'next_uor': 'Middle Office Operations',
    },
    {
        'id': 'STO.3',
        'title': 'Middle Office Operations',
        'objective': 'Explain the role of the middle office in risk management, trade validation, and P&L oversight.',
        'competency': 'Middle Office Function',
        'subcompetencies': [
            'SC3.1 Describe the risk management function of the middle office',
            'SC3.2 Explain P&L attribution and independent price verification (IPV)',
            'SC3.3 Identify the middle office role in trade confirmation and exception management',
            'SC3.4 Explain collateral management and margin monitoring',
        ],
        'ears': 'Describe · Explain · Identify',
        'timeline': '11 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Describe & Explain',
                'scenes': [
                    {
                        'num': 1, 'char': 'ARIA',
                        'visual': 'ARIA at desk; middle office as "referee" graphic — standing between front office and back office with a rulebook',
                        'onscreen': 'The Middle Office\nThe governance layer of the STO\nRisk · P&L · Compliance · Collateral\n"Trust but verify — then verify again"',
                        'vo': 'The middle office is the institutional referee. It sits between the revenue-generating front office and the settlement-focused back office, and its job is to ensure that what the front office says it has done is accurate, within limits, and properly valued. The middle office does not execute trades. It does not settle them. It validates them, monitors the risk they create, independently verifies their value, and escalates when something is wrong. In many firms, the middle office reports into risk or finance rather than the trading desk — precisely to preserve independence.',
                    },
                    {
                        'num': 2, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; risk limit framework — trading limits, VaR limits, counterparty limits, concentration limits',
                        'onscreen': 'Risk Management in the Middle Office\nTrading Limits: position size, delta, DV01\nVaR Limit: Value at Risk — max expected loss at 95%/99%\nCounterparty Limits: cap exposure to any single firm\nConcentration Limits: sector or asset class caps',
                        'vo': 'The middle office owns and enforces the risk limit framework. Every trading desk operates within a set of limits approved by the risk committee: position limits in number of shares or notional, Greek limits for options books like delta and gamma, DV01 limits for fixed income desks showing sensitivity to a one basis point rate move, and VaR limits defining the maximum expected daily loss at a given confidence level. The middle office monitors these in real time. Breach a limit — even by one share — and an immediate escalation is required, no matter how experienced the trader.',
                    },
                    {
                        'num': 3, 'char': 'ARIA',
                        'visual': 'ARIA at screen; IPV process graphic — front office price vs independent source (Bloomberg/Reuters) comparison',
                        'onscreen': 'Independent Price Verification (IPV)\nFront office marks positions using their own models/sources\nMiddle office verifies using independent price sources\nMaterial differences → IPV reserve → P&L haircut\nPrevents traders inflating their own book values',
                        'vo': 'Independent Price Verification is one of the most important middle office controls. Every position in the bank\'s books must be valued using prices from independent sources — typically Bloomberg, Reuters, or third-party pricing services — not the trader\'s own models. If the trader marks a structured bond at 102 but the independent source says 99, the middle office applies an IPV reserve: a deduction from the reported P&L reflecting the uncertainty. This prevents the scenario where traders inflate their book values to hit bonus targets — a mechanism that contributed to several high-profile trading scandals.',
                    },
                    {
                        'num': 4, 'char': 'RYO',
                        'visual': 'RYO to camera; P&L attribution waterfall graphic — breaking down total P&L into market moves, carry, new trades',
                        'onscreen': 'P&L Attribution\nBreaks down daily P&L into its sources:\n• Market move: delta P&L from price changes\n• Carry / theta: time value decay\n• New trades: P&L from trades done today\n• Unexplained: triggers investigation',
                        'vo': 'P&L attribution is the process of explaining every pound or rupee of daily profit and loss. The middle office decomposes the total P&L into its components: how much came from market price movements — the delta P&L; how much from the passage of time — carry or theta decay for options positions; how much from trades executed today; and how much is unexplained. Unexplained P&L is a red flag — it indicates a pricing error, a missing position, or a system problem. Zero unexplained P&L is the target. Any residual triggers an investigation before the books close.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Identify & Explain (Confirmation & Collateral)',
                'scenes': [
                    {
                        'num': 5, 'char': 'ARIA',
                        'visual': 'ARIA at desk; trade confirmation process graphic — affirmation, confirmation, matching via CMS',
                        'onscreen': 'Trade Confirmation Process\nAffirmation: front office agrees to key trade terms\nConfirmation: formal legal document exchanged with counterparty\nMatching: counterparty\'s confirmation matches ours\nUnmatched confirms → follow-up within T+1',
                        'vo': 'Once a trade is executed, confirmation must follow — and this is a middle office responsibility. Affirmation is the first step: both sides agree on the economic terms — instrument, direction, quantity, price, settlement date. Then formal confirmations are exchanged — for OTC trades, typically via electronic platforms like DTCC TIW or Markit MatchIT. If both parties\' confirmations match, the trade is confirmed and flows to settlement. Unmatched confirmations must be resolved by T+1 — failing to confirm in time is a compliance breach and a settlement risk.',
                    },
                    {
                        'num': 6, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; collateral management workflow — variation margin call, collateral transfer, substitution',
                        'onscreen': 'Collateral Management\nFor OTC derivatives under ISDA/CSA\nVariation Margin: daily cash or securities posted for MTM moves\nInitial Margin: upfront buffer for potential future exposure\nCollateral calls issued and received daily by middle office',
                        'vo': 'For OTC derivatives, the middle office manages collateral — the financial buffer posted between counterparties to cover mark-to-market exposure. Under ISDA Credit Support Annexes, parties exchange variation margin daily: if your swap is in-the-money by Rs 10 crore, your counterparty posts Rs 10 crore in cash or eligible securities to you. The middle office calculates the exposure, issues the margin call, monitors receipt, and handles disputes. Initial margin — a larger upfront buffer — is also required under post-2008 regulations for non-centrally-cleared derivatives.',
                    },
                    {
                        'num': 7, 'char': 'ARIA',
                        'visual': 'ARIA at screen; exception management dashboard — aged confirmations, limit breaches, IPV differences highlighted in red',
                        'onscreen': 'Exception Management\nMiddle office runs daily exception reports\n• Unconfirmed trades > T+1\n• Limit breaches (any size)\n• IPV differences > threshold\n• Collateral disputes\nAll exceptions escalated and time-stamped',
                        'vo': 'The middle office does not just verify — it manages exceptions. Every morning, exception reports flag items that require action: trades that have not been confirmed within the required timeframe, limit breaches from the previous day, IPV differences above the materiality threshold, and collateral disputes where counterparties disagree on the net exposure. Each exception is assigned to an owner, given a deadline, and tracked to resolution. Aged exceptions — items open for more than three days without resolution — are escalated to senior management. The exception log is evidence of operational control.',
                    },
                    {
                        'num': 8, 'char': 'RYO',
                        'visual': 'RYO smiling; summary card; "Next: Back Office Operations" teaser',
                        'onscreen': 'Middle Office Summary\n• Risk management: enforces limits in real time\n• IPV: independent valuation prevents P&L inflation\n• P&L attribution: explains every pound of daily P&L\n• Collateral management: daily margin calls under ISDA/CSA',
                        'vo': 'The middle office is the control layer that sits between revenue generation and settlement. It verifies positions, enforces risk limits, validates valuations, and manages the collateral engine that keeps OTC derivative books solvent. Without a functioning middle office, front office risk would go unchecked and settlement failures would multiply. In the next session, we move to the back office — where trades become real: where securities move and cash settles.',
                    },
                ],
            },
        ],
        'next_uor': 'Back Office Operations',
    },
    {
        'id': 'STO.4',
        'title': 'Back Office Operations',
        'objective': 'Describe the back office functions of settlement, custody, and reconciliation, and explain how settlement failures are managed.',
        'competency': 'Back Office & Settlement Function',
        'subcompetencies': [
            'SC4.1 Describe the settlement process and the roles of custodians and CSDs',
            'SC4.2 Explain the concept of DVP (Delivery versus Payment)',
            'SC4.3 Identify common causes of settlement failures and the remediation process',
            'SC4.4 Describe reconciliation between internal books and external statements',
        ],
        'ears': 'Describe · Explain · Identify',
        'timeline': '12 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Describe & Explain (Settlement)',
                'scenes': [
                    {
                        'num': 1, 'char': 'ARIA',
                        'visual': 'ARIA at desk; settlement infrastructure map — buyer, seller, custodians, CSD (NSDL/CDSL/Euroclear), CCP',
                        'onscreen': 'Settlement Infrastructure\nBuyer and Seller → instruct their custodians\nCustodians → instruct the CSD\nCSD: Central Securities Depository (NSDL/CDSL/Euroclear)\nDVP: Securities and cash move simultaneously',
                        'vo': 'Settlement is the moment a trade becomes real. The buyer pays and receives securities; the seller delivers securities and receives cash. In modern markets, this does not happen directly between buyer and seller — it happens through a chain of intermediaries. Each party instructs its custodian. The custodians communicate with the Central Securities Depository — NSDL or CDSL in India, Euroclear or DTC internationally. The CSD records the change of beneficial ownership and simultaneously moves the cash. This simultaneous exchange — Delivery versus Payment, or DVP — eliminates the risk that one side delivers without receiving.',
                    },
                    {
                        'num': 2, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; DVP mechanism step-by-step — securities leg and cash leg linked in CSD',
                        'onscreen': 'Delivery versus Payment (DVP)\nStep 1: Both instructions received by CSD\nStep 2: Securities leg and cash leg matched\nStep 3: Simultaneous transfer — atomic settlement\nNo securities without cash; no cash without securities',
                        'vo': 'DVP is the cornerstone of safe settlement. Without it, the settlement would be sequential — one party moves first and trusts the other to follow. This creates Herstatt Risk: named after a German bank that failed in 1974 after receiving currency from counterparties but before making its own payments. DVP solves this by making settlement atomic — securities and cash transfer simultaneously. If either leg fails, neither moves. The CSD holds both instructions and only executes when both sides are confirmed. DVP is now the global standard, mandated by regulators worldwide.',
                    },
                    {
                        'num': 3, 'char': 'ARIA',
                        'visual': 'ARIA at screen; settlement cycle graphic — T+1 vs T+2 timeline, India vs US vs UK comparison',
                        'onscreen': 'Settlement Cycles\nIndia (equities): T+1 (since 2023 — world\'s fastest)\nUS equities: T+1 (since May 2024)\nEurope (equities): T+2\nBonds: typically T+1 or T+2 depending on market\nT = Trade Date | +1 = one business day after',
                        'vo': 'Settlement cycles define how long after a trade is executed the actual exchange of securities and cash must occur. India moved to T+1 equity settlement in January 2023, making it one of the fastest in the world — and a source of national pride for SEBI. The US followed with T+1 in May 2024. Europe operates on T+2 for most equities. A shorter settlement cycle reduces counterparty risk — the window in which one party can default between trade date and settlement — but demands faster operations and more efficient collateral management.',
                    },
                    {
                        'num': 4, 'char': 'RYO',
                        'visual': 'RYO serious expression; settlement fail causes graphic — SSI error, insufficient securities, cash shortfall, system cutoff missed',
                        'onscreen': 'Why Trades Fail to Settle\n1. Incorrect SSI (Standard Settlement Instructions)\n2. Insufficient securities in custodian account\n3. Cash shortfall on value date\n4. System or communication failure\n5. Trade dispute — terms not agreed\nCSRD/CSDR penalties in Europe; SEBI fines in India',
                        'vo': 'Settlement failures — fails — are costly and reputationally damaging. The most common cause is incorrect Standard Settlement Instructions: the wrong account number, wrong depository code, wrong BIC. If your instructions route the securities to the wrong custodian, they will not arrive where expected. Other causes: the seller does not actually hold the securities they agreed to deliver — a short-selling oversight. The buyer does not have sufficient cash. A system cutoff time is missed. Under CSDR in Europe and SEBI regulations in India, firms face financial penalties for settlement failures. Buy-ins — where the exchange buys the securities from a third party at the failing party\'s cost — are a last resort.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Identify & Describe (Custody & Recon)',
                'scenes': [
                    {
                        'num': 5, 'char': 'ARIA',
                        'visual': 'ARIA at desk; custody service model — global custodian, sub-custodian network, asset servicing roles',
                        'onscreen': 'Custody Services\nGlobal Custodian: holds and safeguards client assets\nSub-custodians: local market specialists\nAsset Servicing: dividends, corporate actions, income collection\nSafe-keeping: separate from firm\'s own assets (client money rules)',
                        'vo': 'Custody is the safe-keeping of securities on behalf of clients. A global custodian — think BNY Mellon, State Street, or Citi — holds securities in segregated accounts, completely separate from the firm\'s own assets. For markets where the custodian has no local presence, they appoint sub-custodians — local banks who are members of the domestic CSD. Beyond safe-keeping, custodians provide asset servicing: collecting dividends and coupon payments, processing corporate actions like rights issues and stock splits, and providing securities lending services. For investment managers, the custodian is the operational backbone of their entire portfolio.',
                    },
                    {
                        'num': 6, 'char': 'RYO',
                        'visual': 'RYO at whiteboard; reconciliation process graphic — internal records vs custodian statement vs CSD records',
                        'onscreen': 'Reconciliation in the Back Office\nThree-way reconciliation: Internal books ↔ Custodian statement ↔ CSD\nDaily for cash and securities positions\nBreaks: differences that must be investigated and resolved\nAged breaks > 3 days = escalation to management',
                        'vo': 'Reconciliation is the back office\'s most critical control. Every day, operations teams compare three sets of records: the firm\'s internal position records — what the trade capture system says you hold; the custodian\'s statement — what they record as held in your account; and the CSD\'s records — the official book of record. Any difference is called a break. Breaks can arise from timing differences — a trade executed today that settles tomorrow — or from genuine errors. Timing breaks resolve automatically. Error breaks require investigation. Any break aged more than three business days must be escalated to senior management and reported in the daily exception report.',
                    },
                    {
                        'num': 7, 'char': 'ARIA',
                        'visual': 'ARIA at screen; fails management workflow — aged fails report, buy-in notice, close-out process',
                        'onscreen': 'Fails Management\nDay 1: settlement instruction sent\nDay 2 (T+1 fail): first chase to counterparty\nDay 4+: escalation; potential buy-in notice issued\nBuy-in: exchange purchases securities at failing party\'s cost\nPenalties: cash compensation under CSDR/SEBI',
                        'vo': 'A fail is a live problem that compounds daily. On the day of the fail, the back office sends a chase to the failing counterparty — or receives one. If the fail persists into the next day, escalation follows: the operations manager is notified, the front office is informed, and the counterparty is pressed. After a defined number of days — typically four in Europe under CSDR — a buy-in notice can be issued. A buy-in is an extreme measure: the exchange or clearing agent purchases the securities from the open market and charges the failing party for any price difference plus costs. It is expensive, embarrassing, and escalates to senior management.',
                    },
                    {
                        'num': 8, 'char': 'RYO',
                        'visual': 'RYO smiling; back office summary card; "Next: Reference Data Management" teaser',
                        'onscreen': 'Back Office Summary\n• Settlement: DVP through CSD — simultaneous, atomic\n• T+1 settlement in India and US (2024)\n• Custody: safe-keeping + asset servicing\n• Reconciliation: three-way daily — breaks must be resolved',
                        'vo': 'The back office is the operational foundation of the STO. Settlement — the movement of securities and cash — happens through custodians and CSDs using the DVP mechanism. India operates T+1, one of the world\'s most efficient settlement regimes. Custody ensures assets are safely held and income is collected. And reconciliation is the daily discipline that ensures the firm\'s books match reality. Next: we look at the critical infrastructure that makes all of this possible — Reference Data Management.',
                    },
                ],
            },
        ],
        'next_uor': 'Reference Data Management (RDM)',
    },
    {
        'id': 'STO.5',
        'title': 'Reference Data Management (RDM)',
        'objective': 'Explain the role of reference data in the STO and describe the key types of reference data and their management.',
        'competency': 'Reference Data Fundamentals',
        'subcompetencies': [
            'SC5.1 Define reference data and distinguish it from market/transactional data',
            'SC5.2 Identify key types: security master, counterparty data, SSI, corporate actions',
            'SC5.3 Explain the data quality impact on trade processing and settlement',
            'SC5.4 Describe the data governance and reference data management lifecycle',
        ],
        'ears': 'Define · Identify · Explain · Describe',
        'timeline': '12 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Define & Identify',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at desk; iceberg graphic — market data visible above the surface, reference data as the massive invisible base',
                        'onscreen': 'Reference Data\nThe invisible foundation of every trade\nNot market prices — the STATIC information about instruments and parties\nIf reference data is wrong, everything downstream is wrong',
                        'vo': 'Reference data is the invisible foundation of capital markets. When you think of financial data, you think of prices — Bloomberg tickers, live bid-ask spreads, real-time charts. That is market data. Reference data is everything else: the static, foundational information about financial instruments, counterparties, and legal entities that every trade relies on. The ISIN code that identifies the security. The LEI that identifies the counterparty. The settlement instructions that tell the custodian where to deliver securities. Get reference data wrong, and everything downstream fails — confirmations break, settlements fail, and P&L is misbookings.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; four boxes: Security Master, Counterparty/Entity Data, SSI, Corporate Actions — with examples in each',
                        'onscreen': 'Key Types of Reference Data\n1. Security Master: ISIN, CUSIP, ticker, asset class, currency, exchange\n2. Counterparty Data: LEI, BIC, legal name, KYC status\n3. SSI (Standard Settlement Instructions): custodian, account, depository\n4. Corporate Actions: dividends, splits, mergers, rights issues',
                        'vo': 'Reference data has four major categories. First, the security master: a record for every financial instrument — its ISIN or CUSIP identifier, its name, asset class, currency of denomination, primary exchange, tick size. Second, counterparty and entity data: Legal Entity Identifiers, SWIFT BIC codes, legal names, credit ratings, KYC and AML status. Third, Standard Settlement Instructions: the specific custodian account and depository details for each counterparty — the routing instructions for delivering securities. Fourth, corporate actions data: upcoming dividends, stock splits, rights issues, mergers — events that affect the value and quantity of securities held.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'RYO serious; error cascade graphic — bad SSI → settlement fail → failed delivery → client complaint → regulatory fine',
                        'onscreen': 'Data Quality = Operational Quality\nBad ISIN → trade booked to wrong security\nBad SSI → settlement instruction sent to wrong custodian\nBad corporate action → wrong dividend credited\nEach error cascades — cost grows at every step',
                        'vo': 'The quality of reference data directly determines the quality of operations. A bad ISIN — say, using an old or incorrect identifier — means a trade is booked against the wrong security. When settlement is instructed, it goes to the wrong counterparty or fails entirely. A missing SSI means no settlement instruction can be generated — the trade sits unprocessed until someone manually finds the correct routing. A stale corporate action record means dividends are credited on the wrong date, or a client receives the wrong number of shares after a stock split. Each error cascades: the later it is caught, the more expensive the fix.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA at screen; reference data sources — Bloomberg, Reuters/Refinitiv, DTCC, SWIFT, local exchange feeds',
                        'onscreen': 'Sources of Reference Data\nBloomberg / Refinitiv: security master, pricing reference\nDTCC / Omgeo: trade matching and SSI golden copy\nSWIFT BIC Directory: financial institution identifiers\nGLEIF: Legal Entity Identifiers (LEI)\nExchanges & CSDs: official instrument data\nChallenge: multiple sources, often contradictory',
                        'vo': 'Reference data comes from multiple external sources — and they do not always agree. Bloomberg and Refinitiv provide security master data, but their ISIN mappings occasionally differ. DTCC and Omgeo provide a "golden copy" of SSI data for DTCC member firms. SWIFT operates the BIC directory. The GLEIF manages the global LEI database. Local exchanges and CSDs publish their own instrument data. The data management challenge: ingest all these sources, resolve conflicts, apply a hierarchy of authority, and maintain a single clean golden record that every downstream system uses. This is the core problem that Reference Data Management exists to solve.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Explain & Describe (Governance & Lifecycle)',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO at desk; golden copy concept graphic — multiple raw data feeds → normalisation → validation → golden record → downstream systems',
                        'onscreen': 'The Golden Copy Concept\nSingle, authoritative reference data record\nAll downstream systems consume from one source\nEliminates version drift and conflicting records\nMaintained by a dedicated Reference Data team',
                        'vo': 'The golden copy is the central goal of reference data management. Rather than allowing each system to hold its own version of an SSI or ISIN — which inevitably drift apart — the firm maintains one authoritative, validated record for each entity: one ISIN record, one SSI per counterparty per currency, one LEI. Every downstream system — OMS, risk, settlement, reporting — reads from this single source. This eliminates the failure mode where the front office books a trade using one ISIN version while the back office settles using a slightly different one. One record. One truth. Always.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; reference data lifecycle — Create, Validate, Approve, Distribute, Amend, Retire',
                        'onscreen': 'Reference Data Lifecycle\n1. Create: new instrument or counterparty on-boarded\n2. Validate: checked against source data\n3. Approve: four-eyes sign-off before use\n4. Distribute: fed to all consuming systems\n5. Amend: changes controlled via change request\n6. Retire: deactivated when no longer needed',
                        'vo': 'Reference data is managed through a structured lifecycle. When a new security is issued or a new counterparty is on-boarded, a record is created in the reference data system. It is validated against authoritative external sources — Bloomberg, CSD feeds. A four-eyes approval ensures no single person can introduce bad data. Once approved, it is distributed to all consuming systems via automated feeds. Any amendments — a change of counterparty name, an updated SSI — go through the same controlled change request process. Retired records are deactivated, not deleted, preserving an audit trail for historical trades.',
                    },
                    {
                        'num': 7, 'char': 'RYO',
                        'visual': 'RYO at screen; SSI management deep-dive — domestic vs international settlement, SWIFT network codes, custodian account hierarchy',
                        'onscreen': 'SSI Deep Dive\nSSI = Standard Settlement Instructions\nSpecifies: custodian, account, depository, SWIFT codes\nDifferent SSI per counterparty per currency per asset class\nDTCC ALERT/Omgeo: industry SSI database\nStale SSI = the #1 cause of preventable settlement fails',
                        'vo': 'Standard Settlement Instructions deserve special attention because stale SSIs are the single most common cause of preventable settlement failures. An SSI specifies exactly where securities or cash should be delivered for a given counterparty — which custodian bank, which account number, which CSD participant code. A large bank may have hundreds of counterparties, each with different SSIs for equities, bonds, FX, and derivatives. The industry solution is DTCC ALERT or Omgeo — centrally managed SSI databases where counterparties publish their own official settlement instructions. Using an industry-maintained SSI database dramatically reduces fail rates.',
                    },
                    {
                        'num': 8, 'char': 'ARIA',
                        'visual': 'ARIA smiling; RDM summary card; "Next: Corporate Actions" teaser',
                        'onscreen': 'RDM Summary\n• Reference data: static foundation enabling every trade\n• Four types: Security Master, Entity Data, SSI, Corporate Actions\n• Golden copy: single authoritative record consumed by all systems\n• Lifecycle: Create → Validate → Approve → Distribute → Amend → Retire',
                        'vo': 'Reference data is the unglamorous but utterly essential foundation of capital markets. Without correct ISINs, trades are mis-booked. Without correct SSIs, trades fail to settle. Without correct corporate actions data, dividends and elections are wrong. The golden copy concept ensures every system reads from one authoritative source. The lifecycle ensures data quality is maintained from creation to retirement. The RDM team may not be as visible as the trading desk, but when reference data fails, the whole STO feels it. Next: corporate actions — the events that change the nature of securities you hold.',
                    },
                ],
            },
        ],
        'next_uor': 'Corporate Actions Processing',
    },
    {
        'id': 'STO.6',
        'title': 'Corporate Actions Processing',
        'objective': 'Identify the main types of corporate actions and describe the operational process for managing mandatory and voluntary events.',
        'competency': 'Corporate Actions Awareness',
        'subcompetencies': [
            'SC6.1 Define corporate actions and distinguish mandatory from voluntary events',
            'SC6.2 Identify key corporate action types: dividends, rights issues, stock splits, mergers',
            'SC6.3 Describe the operational process for a mandatory corporate action',
            'SC6.4 Explain the risks of failing to process corporate actions correctly',
        ],
        'ears': 'Define · Identify · Describe · Explain',
        'timeline': '10 mins',
        'videos': [
            {
                'label': 'Video 1 (Scenes 1–4) — EAR: Define & Identify',
                'scenes': [
                    {
                        'num': 1, 'char': 'RYO',
                        'visual': 'RYO at desk; corporate actions calendar graphic — dividends, splits, mergers, rights on a timeline',
                        'onscreen': 'Corporate Actions\nEvents that change the economic value or structure of a security\nMandatory: happen automatically — shareholder has no choice\nVoluntary: shareholder must elect an option\nOperations must process every event for every holding',
                        'vo': 'Corporate actions are events declared by a company that affect the holders of its securities. A dividend payment adds cash to your account. A stock split doubles your share count and halves the price per share. A rights issue offers you the right to buy additional shares at a discount. A merger means your shares in Company A are exchanged for shares in Company B. These events are not optional to manage — if the back office misses a corporate action, clients lose income, positions are wrongly stated, and regulatory reports are incorrect. For a custody bank with hundreds of thousands of client positions, corporate actions processing is a round-the-clock operation.',
                    },
                    {
                        'num': 2, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; mandatory vs voluntary matrix — with examples in each quadrant',
                        'onscreen': 'Mandatory vs Voluntary Corporate Actions\n\nMandatory (automatic):\n• Cash dividends · Stock splits · Mergers (no election)\n• Bankruptcy distributions · Name changes\n\nVoluntary (election required):\n• Rights issues (take up or sell nil-paid rights)\n• Tender offers · Dividend reinvestment (DRIP)\n• Conversion of convertible bonds',
                        'vo': 'The critical distinction is mandatory versus voluntary. Mandatory events happen automatically: a cash dividend is credited to your account on the payment date whether you do anything or not. A stock split changes your share count automatically. A mandatory exchange merger converts your old shares to new shares automatically on the effective date. Voluntary events require the holder to make a decision: in a rights issue, you must decide whether to exercise your rights, sell them, or let them lapse. In a tender offer, you decide whether to tender your shares at the offered price. Each voluntary event has a deadline — miss it, and the default option applies, which may not be in the client\'s best interest.',
                    },
                    {
                        'num': 3, 'char': 'RYO',
                        'visual': 'RYO at screen; dividend key dates calendar — Ex-Dividend Date, Record Date, Payment Date with T+1 settlement overlay',
                        'onscreen': 'Dividend Key Dates\nDeclaration Date: company announces dividend\nEx-Dividend Date: buy before this → entitled to dividend\nRecord Date: official holder-of-record list compiled\nPayment Date: cash credited to accounts\nWith T+1 settlement: ex-date = record date – 1 business day',
                        'vo': 'Dividends are the simplest corporate action — but they require precise date management. The declaration date is when the company announces the dividend amount and key dates. The ex-dividend date is the critical date: if you buy the share before ex-date, you are entitled to the dividend; if you buy on or after ex-date, the seller keeps it. The record date is when the company compiles its shareholder register — whoever is registered as a holder on record date receives the dividend. With T+1 settlement in India and the US, the ex-dividend date is now one business day before the record date. Payment date is when cash actually arrives in accounts.',
                    },
                    {
                        'num': 4, 'char': 'ARIA',
                        'visual': 'ARIA to camera; stock split mechanics — 2-for-1 split: 100 shares at Rs 500 become 200 shares at Rs 250',
                        'onscreen': 'Stock Splits and Bonus Issues\nStock Split (2-for-1): double shares, halve price\nBonus Issue (1:1): extra shares free from reserves\nEffect: no change in total market cap\nOperations: update all position records on ex-date\nRecords in securities master must update simultaneously',
                        'vo': 'Stock splits and bonus issues change the number of shares without changing the total value. In a 2-for-1 split, every shareholder receives one extra share for each share held — 100 shares at Rs 500 becomes 200 shares at Rs 250. The market cap is unchanged. A bonus issue is similar but funded from the company\'s reserves. For operations, the challenge is that every position record — in the trade capture system, the risk system, the custody system, and the client\'s portfolio management system — must be updated simultaneously on the ex-date. Missing an update in one system creates a reconciliation break. Given the volume of positions at a large custodian, this is automated but monitored carefully.',
                    },
                ],
            },
            {
                'label': 'Video 2 (Scenes 5–8) — EAR: Describe & Explain (Operations Process)',
                'scenes': [
                    {
                        'num': 5, 'char': 'RYO',
                        'visual': 'RYO at desk; rights issue operational flow — notification, election window, sub-underwriting, nil-paid rights trading',
                        'onscreen': 'Rights Issue — Voluntary Corporate Action\nCompany offers existing shareholders new shares at a discount\nShareholders elect: Exercise · Sell nil-paid rights · Lapse\nElection deadline: typically 2–3 weeks from announcement\nBack office must: notify clients, receive elections, submit to registrar',
                        'vo': 'A rights issue is where operations complexity peaks for voluntary events. The company announces a rights issue: for every 5 shares held, shareholders can buy 1 new share at, say, a 20% discount to the market price. The back office must: first, identify all clients holding the relevant shares on record date; second, issue notification letters or electronic messages to each client explaining the terms and deadline; third, receive and record each client\'s election; fourth, aggregate all elections and submit them to the company\'s registrar by the deadline. Clients who do not respond default to their stated default option — typically lapse. Missed elections are client money issues.',
                    },
                    {
                        'num': 6, 'char': 'ARIA',
                        'visual': 'ARIA at whiteboard; merger/takeover flow — offer, terms, election, exchange ratio, effective date',
                        'onscreen': 'Mergers & Acquisitions — Corporate Action\nCash takeover: shareholder receives cash per share\nShare-for-share: shareholder receives acquirer shares\nCash + shares: mixed consideration\nOperations: receive elections, process on effective date\nPosition records replaced — old security retired, new security created',
                        'vo': 'A merger or acquisition is the most complex corporate action operationally. In a cash takeover, shareholders receive a fixed cash amount per share — straightforward, but high volume. In a share-for-share exchange, each share of the target is exchanged for a specified number of acquirer shares — the exchange ratio. In a mixed offer, shareholders receive a combination. Operations must manage the election process, verify entitlements, and on the effective date: debit all positions in the target company and credit either cash or acquirer shares. The old security is retired in the security master. The new positions must reconcile with what the registrar has paid. One error at scale affects thousands of accounts.',
                    },
                    {
                        'num': 7, 'char': 'ARIA',
                        'visual': 'ARIA serious; corporate action failure scenarios — missed election deadline, wrong entitlement calculation, FX conversion error',
                        'onscreen': 'Risks in Corporate Actions Processing\n• Missed election deadline: client loses entitlement value\n• Wrong entitlement: incorrect dividend/share allocation\n• FX risk: foreign dividends credited in wrong currency\n• Timing mismatch: settlement of entitlement vs. position\n• Regulatory reporting: corporate events must be reported accurately',
                        'vo': 'Corporate action errors are particularly painful because they directly affect client money. Missing an election deadline means a client does not receive a dividend reinvestment or rights allocation they were entitled to — and the firm must make them whole. Wrong entitlement calculations — say, crediting a dividend on the wrong quantity because a position was updated late — create P&L discrepancies that must be manually investigated and corrected. For foreign dividends, FX conversion errors can credit clients in the wrong currency. And all corporate action events must be accurately reflected in regulatory reports under MiFID II and SEBI reporting frameworks.',
                    },
                    {
                        'num': 8, 'char': 'RYO',
                        'visual': 'RYO smiling; module complete card — "STO & RDM — Module 8 Complete" with all six UORs listed',
                        'onscreen': 'Module STO Complete\n• STO structure: Front, Middle, Back Office\n• Settlement: DVP, T+1, custodians, CSDs\n• Reference Data: golden copy, SSI, security master\n• Corporate Actions: mandatory and voluntary events\nAll roads in capital markets operations run through the STO.',
                        'vo': 'You have now completed the Securities Trading Organisation module. From the high level — the three-division structure and Chinese Wall — through to the operational detail of front office execution, middle office risk governance, back office settlement and custody, reference data management, and corporate actions processing. Every trade that is executed, confirmed, settled, and reported runs through this machinery. Understanding the STO is understanding how capital markets actually work — not just the instruments, but the operational infrastructure that makes those instruments real. Congratulations on completing Module 8.',
                    },
                ],
            },
        ],
        'next_uor': None,
    },
]

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    elif level == 2:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)
    return p

def add_mapping_table(doc, uor):
    p = doc.add_paragraph()
    run = p.add_run('🔒  Internal Mapping Table')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    widths = [Inches(0.9), Inches(1.4), Inches(2.2), Inches(1.0), Inches(1.8)]
    headers = ['UOR ID', 'UOR Title', 'Learning Objectives', 'EARs', 'Sub-competencies']
    hrow = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = hrow.cells[i]
        cell.width = w
        set_cell_bg(cell, '1E3A5F')
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(hdr)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(9)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    drow = table.rows[1]
    vals = [uor['id'], uor['title'], uor['objective'], uor['ears'], '\n'.join(uor['subcompetencies'])]
    for i, (val, w) in enumerate(zip(vals, widths)):
        cell = drow.cells[i]
        cell.width = w
        set_cell_bg(cell, 'EEF4FB')
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(val)
        run2.font.size = Pt(9)

    doc.add_paragraph()

def add_script_table(doc, video):
    p = doc.add_paragraph()
    run = p.add_run(f'🎬  Production Script — {video["label"]}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    cols = ['Scene #', 'Character', 'Visual Cue', 'On-Screen Text', 'Voice Over']
    col_widths = [Inches(0.5), Inches(0.7), Inches(1.5), Inches(1.6), Inches(3.0)]
    table = doc.add_table(rows=len(video['scenes']) + 1, cols=5)
    table.style = 'Table Grid'

    hrow = table.rows[0]
    for i, (col, w) in enumerate(zip(cols, col_widths)):
        cell = hrow.cells[i]
        cell.width = w
        set_cell_bg(cell, '005F5F')
        p2 = cell.paragraphs[0]
        run2 = p2.add_run(col)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(9)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    alt = ['FFFFFF', 'F0F7F4']
    for r_idx, scene in enumerate(video['scenes']):
        row = table.rows[r_idx + 1]
        fill = alt[r_idx % 2]
        vals = [str(scene['num']), scene['char'], scene['visual'], scene['onscreen'], scene['vo']]
        for i, (val, w) in enumerate(zip(vals, col_widths)):
            cell = row.cells[i]
            cell.width = w
            set_cell_bg(cell, fill)
            p2 = cell.paragraphs[0]
            run2 = p2.add_run(val)
            run2.font.size = Pt(8.5)
            if i == 1:
                run2.bold = True
                run2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F) if val == 'RYO' else RGBColor(0x00, 0x5F, 0x5F)

    doc.add_paragraph()

def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SECURITIES TRADING ORGANISATION & RDM')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Video Production Document')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('CIBOP Module 8 | Version 2')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('Characters: RYO (Male, Navy Suit, Mentor, Dry Wit) | ARIA (Female, Teal Blazer, Analyst, Sharp Comic Timing)')
    run4.font.size = Pt(10)
    run4.italic = True
    doc.add_paragraph()
    doc.add_page_break()

    # TOC
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run('TABLE OF CONTENTS')
    run_toc.bold = True
    run_toc.font.size = Pt(12)
    run_toc.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    for uor in UORS:
        pi = doc.add_paragraph()
        pi.paragraph_format.left_indent = Inches(0.3)
        ri = pi.add_run(f'{uor["id"]}  {uor["title"]}')
        ri.font.size = Pt(10)
    doc.add_page_break()

    # Per-UOR content
    for idx, uor in enumerate(UORS):
        add_heading(doc, f'{uor["id"]} — {uor["title"]}', level=1)
        ptag = doc.add_paragraph()
        rtag = ptag.add_run(f'EARs: {uor["ears"]}  |  Competency: {uor["competency"]}  |  Timeline: {uor["timeline"]}')
        rtag.font.size = Pt(9)
        rtag.italic = True
        rtag.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()

        add_mapping_table(doc, uor)

        for video in uor['videos']:
            add_script_table(doc, video)

        # Transition card
        nxt = uor['next_uor']
        trans_text = f'End of {uor["title"]} — Next: {nxt}' if nxt else f'End of {uor["title"]} — Module Complete'
        doc.add_paragraph()
        ttable = doc.add_table(rows=1, cols=1)
        ttable.style = 'Table Grid'
        tc = ttable.rows[0].cells[0]
        set_cell_bg(tc, '1E3A5F')
        ptc = tc.paragraphs[0]
        ptc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rtc = ptc.add_run(f'▶  {trans_text}')
        rtc.bold = True
        rtc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rtc.font.size = Pt(11)

        if idx < len(UORS) - 1:
            doc.add_page_break()

    # Footer
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        fr = fp.add_run('CIBOP PPT 08 — Securities Trading Organisation & RDM  |  Confidential — Internal Use Only')
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')
    import os
    print(f'File size: {os.path.getsize(OUTPUT_PATH):,} bytes')

if __name__ == '__main__':
    build_doc()
    print('Video Production Document DONE.')
