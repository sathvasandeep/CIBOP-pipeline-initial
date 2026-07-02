#!/usr/bin/env python3
"""Generate STO & RDM Knowledge Assessment — CIBOP Module 8"""
import subprocess, sys
subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'], check=True)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = '/Users/sandeeppr/CIBOP/PPT 08 Securities/STO_Knowledge_Assessment_v2.docx'

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def q_header(doc, num, qtype, ear, uor_id, video_ref):
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    widths = [0.6, 1.8, 1.8, 3.1]
    vals = [num, qtype, f'EAR: {ear}', f'{uor_id} | {video_ref}']
    fills = ['1E3A5F', '005F5F', '8B6914', 'EEEEEE']
    tcolors = [(255,255,255),(255,255,255),(255,255,255),(30,30,30)]
    for i, (v, w, f, tc) in enumerate(zip(vals, widths, fills, tcolors)):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, f)
        set_col_width(cell, w)
        p = cell.paragraphs[0]
        r = p.add_run(str(v))
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(*tc)
    doc.add_paragraph()

def answer_key_box(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, 'FFF8E1')
    p = cell.paragraphs[0]
    r = p.add_run(f'✓  ANSWER KEY: {text}')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)
    doc.add_paragraph()

# ─────────────────────────────────────────────
# QUESTION BANK — 4 questions per UOR, cycling 5 types
# Types: MCQ INLINE | TRUE/FALSE MULTI | GAP FILL DROPDOWN | MATCHING PAIRS | TABLE SELECTION
# ─────────────────────────────────────────────

QUESTIONS = {

    'STO.1': [
        # Q1 — MCQ INLINE
        {
            'type': 'MCQ INLINE', 'ear': 'Identify', 'video': 'Video 1',
            'question': 'Which of the following BEST describes the primary role of the Back Office in a Securities Trading Organisation?',
            'options': [
                ('A', 'Executing trades on behalf of clients', False),
                ('B', 'Monitoring risk limits and P&L attribution', False),
                ('C', 'Confirming, clearing, and settling trades', True),
                ('D', 'Generating investment research and market analysis', False),
            ],
        },
        # Q2 — TRUE/FALSE MULTI
        {
            'type': 'TRUE/FALSE MULTI', 'ear': 'Define', 'video': 'Video 1',
            'question': 'Indicate whether each statement about the Securities Trading Organisation is TRUE or FALSE.',
            'statements': [
                ('The Chinese Wall prevents the front office from accessing client settlement information held by the back office.', True),
                ('Straight-Through Processing (STP) refers to a trade that requires manual intervention at every stage.', False),
                ('The STP rate is a key operational efficiency metric for the back office.', True),
                ('The front office is responsible for confirming and settling trades after execution.', False),
            ],
        },
        # Q3 — GAP FILL DROPDOWN
        {
            'type': 'GAP FILL DROPDOWN', 'ear': 'Explain', 'video': 'Video 2',
            'sentence': 'In the STO trade lifecycle, after a trade is executed by the front office it flows through [BLANK1], then [BLANK2], and finally to [BLANK3] where securities and cash change hands.',
            'blanks': [
                ('BLANK1', 'trade confirmation', ['trade confirmation', 'portfolio management', 'credit rating', 'marketing approval']),
                ('BLANK2', 'clearing', ['clearing', 'origination', 'underwriting', 'distribution']),
                ('BLANK3', 'settlement', ['settlement', 'execution', 'research', 'structuring']),
            ],
        },
        # Q4 — MATCHING PAIRS
        {
            'type': 'MATCHING PAIRS', 'ear': 'Identify', 'video': 'Video 2',
            'instruction': 'Match each STO system to its primary function.',
            'pairs': [
                ('OMS (Order Management System)', 'Creates, tracks, and routes client orders from receipt to fill'),
                ('EMS (Execution Management System)', 'Connects to exchanges and dark pools for market execution'),
                ('Risk System', 'Monitors live trading positions against approved risk limits'),
                ('Settlement System', 'Instructs custodians and clearinghouses to move securities and cash'),
            ],
        },
    ],

    'STO.2': [
        # Q1 — TABLE SELECTION
        {
            'type': 'TABLE SELECTION', 'ear': 'Identify', 'video': 'Video 1',
            'instruction': 'The table below lists individuals and activities. Tick (✓) each row that represents a Front Office role or responsibility.',
            'rows': [
                ('Equity Trader executing a client order on the NSE', True),
                ('Operations analyst confirming trade details with counterparty', False),
                ('Sales staff managing client relationship and channelling orders to traders', True),
                ('Risk manager independently verifying trader P&L', False),
                ('Research analyst publishing an equity sector report', True),
                ('Reconciliation team comparing custodian statements to internal records', False),
            ],
        },
        # Q2 — MCQ INLINE
        {
            'type': 'MCQ INLINE', 'ear': 'Explain', 'video': 'Video 1',
            'question': 'Under best execution obligations (MiFID II / SEBI), which of the following factors must a firm consider when executing a client order?',
            'options': [
                ('A', 'Only the best available price at the time of execution', False),
                ('B', 'Price, speed of execution, likelihood of execution, size, and cost', True),
                ('C', 'Only the commission rate negotiated with the client', False),
                ('D', 'The firm\'s own proprietary trading interests', False),
            ],
        },
        # Q3 — TRUE/FALSE MULTI
        {
            'type': 'TRUE/FALSE MULTI', 'ear': 'Describe', 'video': 'Video 2',
            'question': 'Indicate whether each statement about front office trading is TRUE or FALSE.',
            'statements': [
                ('Proprietary trading involves the bank executing orders solely on behalf of its clients with no market risk to the bank.', False),
                ('A wrong-side error (buying instead of selling) is considered one of the most costly trade errors in the front office.', True),
                ('The Volcker Rule was introduced to restrict bank proprietary trading activity following the 2008 financial crisis.', True),
                ('Front office traders independently verify and mark their own P&L without middle office involvement.', False),
            ],
        },
        # Q4 — GAP FILL DROPDOWN
        {
            'type': 'GAP FILL DROPDOWN', 'ear': 'Explain', 'video': 'Video 1',
            'sentence': 'In agency trading, the bank executes on behalf of the [BLANK1] and earns a [BLANK2], bearing [BLANK3] market risk on the transaction.',
            'blanks': [
                ('BLANK1', 'client', ['client', 'regulator', 'central bank', 'competitor']),
                ('BLANK2', 'commission', ['commission', 'spread income', 'trading profit', 'margin call']),
                ('BLANK3', 'no', ['no', 'full', 'partial', 'unlimited']),
            ],
        },
    ],

    'STO.3': [
        # Q1 — MATCHING PAIRS
        {
            'type': 'MATCHING PAIRS', 'ear': 'Describe', 'video': 'Video 1',
            'instruction': 'Match each middle office function to its correct description.',
            'pairs': [
                ('Independent Price Verification (IPV)', 'Valuing positions using external sources to validate front office marks'),
                ('P&L Attribution', 'Breaking down daily profit and loss into market moves, carry, and new trades'),
                ('Variation Margin', 'Daily cash or securities posted between OTC derivative counterparties based on MTM moves'),
                ('Exception Management', 'Daily reporting and escalation of unconfirmed trades, limit breaches, and IPV differences'),
            ],
        },
        # Q2 — MCQ INLINE
        {
            'type': 'MCQ INLINE', 'ear': 'Explain', 'video': 'Video 1',
            'question': 'What is the PRIMARY purpose of Independent Price Verification (IPV) in the middle office?',
            'options': [
                ('A', 'To ensure traders submit their orders using the correct OMS workflow', False),
                ('B', 'To confirm that client settlement instructions have been received on time', False),
                ('C', 'To prevent traders from inflating their own book valuations by using independent pricing sources', True),
                ('D', 'To reconcile the firm\'s cash accounts with the custodian\'s daily statement', False),
            ],
        },
        # Q3 — TRUE/FALSE MULTI
        {
            'type': 'TRUE/FALSE MULTI', 'ear': 'Identify', 'video': 'Video 2',
            'question': 'Indicate whether each statement about middle office controls is TRUE or FALSE.',
            'statements': [
                ('An IPV reserve is applied when the front office price and the independent source price differ materially.', True),
                ('Variation margin is only required for exchange-traded derivatives, not OTC instruments.', False),
                ('P&L attribution aims to explain every rupee of daily profit and loss, with zero unexplained residual as the target.', True),
                ('The middle office reports directly into the front office trading desk to ensure close collaboration.', False),
            ],
        },
        # Q4 — GAP FILL DROPDOWN
        {
            'type': 'GAP FILL DROPDOWN', 'ear': 'Explain', 'video': 'Video 2',
            'sentence': 'Under an ISDA [BLANK1], counterparties post [BLANK2] daily to cover mark-to-market exposure, and a larger upfront [BLANK3] buffer for potential future exposure.',
            'blanks': [
                ('BLANK1', 'Credit Support Annex (CSA)', ['Credit Support Annex (CSA)', 'Master Confirmation Agreement', 'ISDA Schedule', 'Global Master Repo Agreement']),
                ('BLANK2', 'variation margin', ['variation margin', 'initial margin', 'settlement proceeds', 'haircut adjustment']),
                ('BLANK3', 'initial margin', ['initial margin', 'variation margin', 'performance bond', 'security deposit']),
            ],
        },
    ],

    'STO.4': [
        # Q1 — TABLE SELECTION
        {
            'type': 'TABLE SELECTION', 'ear': 'Describe', 'video': 'Video 1',
            'instruction': 'The table below lists settlement-related statements. Tick (✓) each statement that is CORRECT.',
            'rows': [
                ('DVP (Delivery versus Payment) ensures securities and cash move simultaneously, eliminating Herstatt Risk.', True),
                ('In India, equity settlement moved to T+1 in January 2023, making it one of the fastest globally.', True),
                ('In a DVP settlement, if the cash leg fails the securities leg still completes automatically.', False),
                ('NSDL and CDSL are India\'s Central Securities Depositories.', True),
                ('A custodian holds securities in commingled accounts alongside the firm\'s own proprietary assets.', False),
            ],
        },
        # Q2 — MCQ INLINE
        {
            'type': 'MCQ INLINE', 'ear': 'Explain', 'video': 'Video 1',
            'question': 'Which of the following is the MOST common cause of preventable settlement failures?',
            'options': [
                ('A', 'Market prices moving against the expected settlement value', False),
                ('B', 'Incorrect Standard Settlement Instructions (SSI) routing the trade to the wrong custodian', True),
                ('C', 'The central securities depository being closed on the value date', False),
                ('D', 'The front office executing at a price different from the client\'s limit', False),
            ],
        },
        # Q3 — MATCHING PAIRS
        {
            'type': 'MATCHING PAIRS', 'ear': 'Identify', 'video': 'Video 2',
            'instruction': 'Match each reconciliation break type to its correct resolution.',
            'pairs': [
                ('Timing break', 'Resolves automatically when the trade settles in the next business day'),
                ('Error break (wrong ISIN booked)', 'Requires manual investigation, trade amendment, and reconfirmation'),
                ('Aged break (>3 days)', 'Must be escalated to senior management and logged in exception report'),
                ('Cash reconciliation break', 'Cross-referenced against bank statements and nostro accounts to identify missing credits'),
            ],
        },
        # Q4 — TRUE/FALSE MULTI
        {
            'type': 'TRUE/FALSE MULTI', 'ear': 'Describe', 'video': 'Video 2',
            'question': 'Indicate whether each statement about custody and settlement is TRUE or FALSE.',
            'statements': [
                ('A global custodian holds client securities in accounts that are legally segregated from the firm\'s own proprietary assets.', True),
                ('A buy-in is a routine settlement mechanism used for all trades regardless of failure status.', False),
                ('Reconciliation compares the firm\'s internal position records against custodian statements and CSD records.', True),
                ('Under CSDR in Europe, firms face no financial penalty for settlement failures below a certain size threshold.', False),
            ],
        },
    ],

    'STO.5': [
        # Q1 — MCQ INLINE
        {
            'type': 'MCQ INLINE', 'ear': 'Define', 'video': 'Video 1',
            'question': 'Which of the following BEST distinguishes reference data from market data?',
            'options': [
                ('A', 'Reference data is real-time pricing; market data is historical pricing', False),
                ('B', 'Reference data is static foundational information about instruments and parties; market data is live transactional pricing', True),
                ('C', 'Reference data is generated by traders; market data is generated by operations teams', False),
                ('D', 'Reference data changes every second; market data changes annually', False),
            ],
        },
        # Q2 — MATCHING PAIRS
        {
            'type': 'MATCHING PAIRS', 'ear': 'Identify', 'video': 'Video 1',
            'instruction': 'Match each reference data type to the correct example.',
            'pairs': [
                ('Security Master', 'ISIN: INE002A01018 | Asset class: Equity | Exchange: NSE | Currency: INR'),
                ('Counterparty / Entity Data', 'LEI: 335800ZVKESJ9EVKH008 | BIC: ICICIINBBXXX | Legal name: ICICI Bank Ltd'),
                ('Standard Settlement Instructions (SSI)', 'Custodian: Deutsche Bank Mumbai | Account: 123456789 | Depository: NSDL'),
                ('Corporate Actions Data', 'Ex-dividend date: 15 July | Record date: 16 July | Payment date: 25 July | Dividend: Rs 18 per share'),
            ],
        },
        # Q3 — TRUE/FALSE MULTI
        {
            'type': 'TRUE/FALSE MULTI', 'ear': 'Explain', 'video': 'Video 2',
            'question': 'Indicate whether each statement about reference data management is TRUE or FALSE.',
            'statements': [
                ('The "golden copy" is a single authoritative reference data record consumed by all downstream systems to prevent version drift.', True),
                ('Reference data records should be deleted immediately when they are no longer active to keep the database clean.', False),
                ('Stale SSI data is considered the single most common cause of preventable settlement failures.', True),
                ('Bloomberg, Refinitiv, and DTCC always publish identical reference data, so no reconciliation between sources is needed.', False),
            ],
        },
        # Q4 — GAP FILL DROPDOWN
        {
            'type': 'GAP FILL DROPDOWN', 'ear': 'Describe', 'video': 'Video 2',
            'sentence': 'The reference data lifecycle begins with [BLANK1] of a new record, followed by [BLANK2] against external authoritative sources, and a [BLANK3] sign-off before the record is distributed to downstream systems.',
            'blanks': [
                ('BLANK1', 'creation', ['creation', 'deletion', 'archiving', 'pricing']),
                ('BLANK2', 'validation', ['validation', 'speculation', 'settlement', 'execution']),
                ('BLANK3', 'four-eyes', ['four-eyes', 'single approver', 'client', 'regulator']),
            ],
        },
    ],

    'STO.6': [
        # Q1 — TRUE/FALSE MULTI
        {
            'type': 'TRUE/FALSE MULTI', 'ear': 'Define', 'video': 'Video 1',
            'question': 'Indicate whether each statement about corporate actions is TRUE or FALSE.',
            'statements': [
                ('A mandatory corporate action requires the shareholder to make an election within a specified deadline.', False),
                ('In a 2-for-1 stock split, the total market capitalisation of the company remains unchanged.', True),
                ('The ex-dividend date is the date on which the company pays the dividend to shareholders of record.', False),
                ('A rights issue is an example of a voluntary corporate action that requires the shareholder to elect an option.', True),
            ],
        },
        # Q2 — MATCHING PAIRS
        {
            'type': 'MATCHING PAIRS', 'ear': 'Identify', 'video': 'Video 1',
            'instruction': 'Match each corporate action key date to its correct definition.',
            'pairs': [
                ('Declaration Date', 'Date on which the company announces the corporate action terms'),
                ('Ex-Dividend Date', 'First date on which a buyer of the shares is NOT entitled to the dividend'),
                ('Record Date', 'Date on which the company compiles the official register of shareholders entitled to the action'),
                ('Payment Date', 'Date on which the dividend cash or entitlement is credited to shareholder accounts'),
            ],
        },
        # Q3 — GAP FILL DROPDOWN
        {
            'type': 'GAP FILL DROPDOWN', 'ear': 'Describe', 'video': 'Video 2',
            'sentence': 'In a rights issue, shareholders who do not respond by the election deadline will have their entitlement treated according to the [BLANK1] option, and the back office must [BLANK2] all clients with eligible holdings on [BLANK3].',
            'blanks': [
                ('BLANK1', 'default', ['default', 'forced exercise', 'open market sale', 'director discretion']),
                ('BLANK2', 'notify', ['notify', 'trade against', 'ignore', 'merge']),
                ('BLANK3', 'record date', ['record date', 'payment date', 'ex-date', 'declaration date']),
            ],
        },
        # Q4 — TABLE SELECTION
        {
            'type': 'TABLE SELECTION', 'ear': 'Explain', 'video': 'Video 2',
            'instruction': 'The table below lists corporate action processing risks. Tick (✓) each item that represents a GENUINE operational risk.',
            'rows': [
                ('Missed election deadline causing a client to lose a rights issue entitlement', True),
                ('A stock split that automatically doubles a client\'s share count without any system update needed', False),
                ('Wrong entitlement calculation crediting a dividend on an incorrect position quantity', True),
                ('An FX conversion error when crediting a foreign dividend in the wrong currency', True),
                ('A cash dividend that is automatically credited, requiring no back office involvement whatsoever', False),
            ],
        },
    ],
}

# ─────────────────────────────────────────────
# BUILD ASSESSMENT DOCUMENT
# ─────────────────────────────────────────────
def render_mcq(doc, q, q_num, uor_id):
    q_header(doc, f'Q{q_num}', 'MCQ INLINE', q['ear'], uor_id, q['video'])
    p = doc.add_paragraph()
    r = p.add_run(q['question'])
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph()
    correct_letter = ''
    for letter, text, correct in q['options']:
        po = doc.add_paragraph()
        po.paragraph_format.left_indent = Inches(0.3)
        ro = po.add_run(f'{letter}.  {text}')
        ro.font.size = Pt(10)
        if correct:
            correct_letter = letter
    doc.add_paragraph()
    answer_key_box(doc, f'Correct answer: {correct_letter}')

def render_tf(doc, q, q_num, uor_id):
    q_header(doc, f'Q{q_num}', 'TRUE/FALSE MULTI', q['ear'], uor_id, q['video'])
    p = doc.add_paragraph()
    r = p.add_run(q['question'])
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph()
    t = doc.add_table(rows=len(q['statements']) + 1, cols=3)
    t.style = 'Table Grid'
    for i, (hdr, w) in enumerate(zip(['#', 'Statement', 'True / False'], [0.3, 5.5, 0.9])):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '1E3A5F')
        set_col_width(cell, w)
        ph = cell.paragraphs[0]
        rh = ph.add_run(hdr)
        rh.bold = True
        rh.font.color.rgb = RGBColor(255, 255, 255)
        rh.font.size = Pt(9)
    answers = []
    for i, (stmt, answer) in enumerate(q['statements']):
        row = t.rows[i + 1]
        fill = 'FFFFFF' if i % 2 == 0 else 'F0F7F4'
        set_cell_bg(row.cells[0], fill)
        set_cell_bg(row.cells[1], fill)
        set_cell_bg(row.cells[2], fill)
        set_col_width(row.cells[0], 0.3)
        set_col_width(row.cells[1], 5.5)
        set_col_width(row.cells[2], 0.9)
        row.cells[0].paragraphs[0].add_run(str(i + 1)).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(stmt).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run('True' if answer else 'False').font.size = Pt(9)
        answers.append(f'{i+1}: {"TRUE" if answer else "FALSE"}')
    doc.add_paragraph()
    answer_key_box(doc, ' | '.join(answers))

def render_gap(doc, q, q_num, uor_id):
    q_header(doc, f'Q{q_num}', 'GAP FILL DROPDOWN', q['ear'], uor_id, q['video'])
    p = doc.add_paragraph()
    r = p.add_run('Complete the sentence by selecting the correct option for each blank.')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    ps = doc.add_paragraph()
    rs = ps.add_run(q['sentence'])
    rs.font.size = Pt(10)
    doc.add_paragraph()
    t = doc.add_table(rows=len(q['blanks']) + 1, cols=3)
    t.style = 'Table Grid'
    for i, (hdr, w) in enumerate(zip(['Blank', 'Options', 'Correct Answer'], [0.7, 4.2, 1.8])):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '005F5F')
        set_col_width(cell, w)
        ph = cell.paragraphs[0]
        rh = ph.add_run(hdr)
        rh.bold = True
        rh.font.color.rgb = RGBColor(255, 255, 255)
        rh.font.size = Pt(9)
    answers = []
    for i, (blank_id, correct, options) in enumerate(q['blanks']):
        row = t.rows[i + 1]
        fill = 'FFFFFF' if i % 2 == 0 else 'EEF4FB'
        set_cell_bg(row.cells[0], fill)
        set_cell_bg(row.cells[1], fill)
        set_cell_bg(row.cells[2], fill)
        set_col_width(row.cells[0], 0.7)
        set_col_width(row.cells[1], 4.2)
        set_col_width(row.cells[2], 1.8)
        row.cells[0].paragraphs[0].add_run(f'[{blank_id}]').font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(' | '.join(options)).font.size = Pt(8.5)
        row.cells[2].paragraphs[0].add_run(correct).font.size = Pt(9)
        answers.append(f'{blank_id}: {correct}')
    doc.add_paragraph()
    answer_key_box(doc, ' | '.join(answers))

def render_matching(doc, q, q_num, uor_id):
    q_header(doc, f'Q{q_num}', 'MATCHING PAIRS', q['ear'], uor_id, q['video'])
    p = doc.add_paragraph()
    r = p.add_run(q['instruction'])
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph()
    t = doc.add_table(rows=len(q['pairs']) + 1, cols=2)
    t.style = 'Table Grid'
    for i, (hdr, w) in enumerate(zip(['Left Column', 'Match to Right Column'], [3.4, 3.4])):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '1E3A5F')
        set_col_width(cell, w)
        ph = cell.paragraphs[0]
        rh = ph.add_run(hdr)
        rh.bold = True
        rh.font.color.rgb = RGBColor(255, 255, 255)
        rh.font.size = Pt(9)
    for i, (left, right) in enumerate(q['pairs']):
        row = t.rows[i + 1]
        fill = 'FFFFFF' if i % 2 == 0 else 'F0F7F4'
        set_cell_bg(row.cells[0], fill)
        set_cell_bg(row.cells[1], fill)
        set_col_width(row.cells[0], 3.4)
        set_col_width(row.cells[1], 3.4)
        row.cells[0].paragraphs[0].add_run(left).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(right).font.size = Pt(9)
    doc.add_paragraph()
    answer_key_box(doc, 'Pairs matched as displayed above (left → right in order).')

def render_table_sel(doc, q, q_num, uor_id):
    q_header(doc, f'Q{q_num}', 'TABLE SELECTION', q['ear'], uor_id, q['video'])
    p = doc.add_paragraph()
    r = p.add_run(q['instruction'])
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph()
    t = doc.add_table(rows=len(q['rows']) + 1, cols=2)
    t.style = 'Table Grid'
    for i, (hdr, w) in enumerate(zip(['Statement', 'Select (✓ if correct)'], [5.8, 1.0])):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, '005F5F')
        set_col_width(cell, w)
        ph = cell.paragraphs[0]
        rh = ph.add_run(hdr)
        rh.bold = True
        rh.font.color.rgb = RGBColor(255, 255, 255)
        rh.font.size = Pt(9)
    correct_rows = []
    for i, (stmt, correct) in enumerate(q['rows']):
        row = t.rows[i + 1]
        fill = 'FFFFFF' if i % 2 == 0 else 'F0F7F4'
        set_cell_bg(row.cells[0], fill)
        set_cell_bg(row.cells[1], fill)
        set_col_width(row.cells[0], 5.8)
        set_col_width(row.cells[1], 1.0)
        row.cells[0].paragraphs[0].add_run(stmt).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run('✓' if correct else '').font.size = Pt(9)
        if correct:
            correct_rows.append(str(i + 1))
    doc.add_paragraph()
    answer_key_box(doc, f'Correct rows: {", ".join(correct_rows)}')

RENDERERS = {
    'MCQ INLINE': render_mcq,
    'TRUE/FALSE MULTI': render_tf,
    'GAP FILL DROPDOWN': render_gap,
    'MATCHING PAIRS': render_matching,
    'TABLE SELECTION': render_table_sel,
}

UOR_TITLES = {
    'STO.1': 'Introduction to the Securities Trading Organisation',
    'STO.2': 'Front Office Operations',
    'STO.3': 'Middle Office Operations',
    'STO.4': 'Back Office Operations',
    'STO.5': 'Reference Data Management (RDM)',
    'STO.6': 'Corporate Actions Processing',
}

def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SECURITIES TRADING ORGANISATION & RDM')
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Knowledge Assessment')
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x00, 0x5F, 0x5F)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('CIBOP Module 8 | Version 2 | 4 Questions per UOR | 5 Question Types')
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)
    doc.add_paragraph()
    doc.add_page_break()

    q_num = 1
    for uor_id, questions in QUESTIONS.items():
        # UOR Section header
        p_hdr = doc.add_paragraph()
        r_hdr = p_hdr.add_run(f'{uor_id} — {UOR_TITLES[uor_id]}')
        r_hdr.bold = True
        r_hdr.font.size = Pt(14)
        r_hdr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        doc.add_paragraph()

        for q in questions:
            renderer = RENDERERS[q['type']]
            renderer(doc, q, q_num, uor_id)
            q_num += 1

        doc.add_page_break()

    # Footer
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        fr = fp.add_run('CIBOP PPT 08 — Securities Trading Organisation & RDM  |  Knowledge Assessment  |  Confidential — Internal Use Only')
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')
    import os
    print(f'File size: {os.path.getsize(OUTPUT_PATH):,} bytes')

if __name__ == '__main__':
    build_doc()
    print('Knowledge Assessment DONE.')
